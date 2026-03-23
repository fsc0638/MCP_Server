"""
LINE Connector — MCP_Server Integration Module
===============================================
Integrates LINE Messaging API into the existing FastAPI + UMA architecture.

Architecture:
  Inbound  → POST /api/line/webhook   (signature verify + parse events)
  Router   → model_router.route_model() (LLM-as-a-Router: nano classifies → tier)
  Session  → core.session.SessionManager  (session_id = "line_{user_id}")
  Outbound → OpenAIAdapter.chat() → reply_message / push_message fallback

Model Routing (省流省錢):
  nano  (200K TPM) → 閒聊 / 打招呼 / 簡單問答
  mini  (200K TPM) → 一般任務 / 翻譯 / 摘要 / 文件分段摘要
  full  (30K  TPM) → 深度分析 / 多工具串接 / 最終合成報告

Design Principle:
  - 立即回覆 200 OK (解決 LINE Webhook 1~2 秒 Timeout 限制)
  - BackgroundTasks 非同步執行 LLM 生成與 Tool Calling
  - 完全重用現有 SessionManager + MEMORY.md 持久化機制
  - LLM-as-a-Router 前置分類，避免所有請求都打高成本模型
  - Token 預估 + 自動降級 (Fallback)
"""
import logging
import os
import threading
from contextlib import contextmanager

from fastapi import APIRouter, Request, BackgroundTasks, HTTPException
import httpx
try:
    import redis
except ImportError:
    redis = None

logger = logging.getLogger("MCP_Server.LINE")
router = APIRouter()

# ── Perpetual Session Date Tracker ───────────────────────────────────────────
# {session_id: YYYY-MM-DD}
_session_days = {}

# ── Lazy-initialized LINE SDK components ──────────────────────────────────────
# 延遲初始化，確保缺少 key 時伺服器仍可啟動（降級模式）
_line_handler = None
_line_api = None
_line_api_blob = None


def _get_line_components():
    """
    Lazily initialize LINE SDK WebhookHandler and MessagingApi.
    Raises KeyError if LINE_CHANNEL_SECRET or LINE_CHANNEL_ACCESS_TOKEN is missing.
    """
    global _line_handler, _line_api, _line_api_blob
    if _line_handler is None:
        from linebot.v3 import WebhookHandler
        from linebot.v3.messaging import Configuration, ApiClient, MessagingApi, MessagingApiBlob

        secret = os.environ.get("LINE_CHANNEL_SECRET", "").strip()
        token = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()

        if not secret or not token:
            raise KeyError(
                "LINE_CHANNEL_SECRET and LINE_CHANNEL_ACCESS_TOKEN must be set in .env"
            )

        cfg = Configuration(access_token=token)
        _api_client = ApiClient(cfg)
        _line_api = MessagingApi(_api_client)
        _line_api_blob = MessagingApiBlob(_api_client)
        _line_handler = WebhookHandler(secret)
        logger.info("[LINE] SDK initialized successfully.")

    return _line_handler, _line_api, _line_api_blob


# ── LINE-specific system prompt (Moved to server/services/runtime.py) ─────────
from server.services.runtime import get_universal_system_prompt

# ── LINE 單則訊息字元上限 ───────────────────────────────────────────────────────
_LINE_MAX_CHARS = 4900


# ── Webhook Endpoint ──────────────────────────────────────────────────────────

@router.post("/api/line/webhook", tags=["Integration"])
async def line_webhook(request: Request, background_tasks: BackgroundTasks):
    """
    LINE Messaging API Webhook 接收端點。

    流程：
    A. 驗證 X-Line-Signature（防偽造請求）
    B. 解析 LINE Events
    C. TextMessage → 丟入 BackgroundTasks（解耦 LLM 延遲）
    D. 立即回覆 200 OK（解決 Timeout 瓶頸）
    """
    # A. 取得並驗證 Signature
    try:
        handler, line_api, line_api_blob = _get_line_components()
    except KeyError as e:
        logger.error(f"[LINE] Missing configuration: {e}")
        raise HTTPException(status_code=500, detail=f"LINE configuration error: {e}")

    signature = request.headers.get("X-Line-Signature", "")
    body_bytes = await request.body()
    body_text = body_bytes.decode("utf-8")

    # B. 解析事件
    from linebot.v3.exceptions import InvalidSignatureError
    from linebot.v3.webhooks import MessageEvent, TextMessageContent

    try:
        events = handler.parser.parse(body_text, signature)
    except InvalidSignatureError:
        logger.warning("[LINE] Webhook rejected — invalid X-Line-Signature")
        raise HTTPException(status_code=400, detail="Invalid signature")
    except Exception as e:
        logger.error(f"[LINE] Event parse error: {e}")
        raise HTTPException(status_code=400, detail=f"Event parse error: {e}")

    # C. 逐一處理 TextMessage / ImageMessage / FileMessage / StickerMessage Event
    from linebot.v3.webhooks import MessageEvent, TextMessageContent, ImageMessageContent, FileMessageContent, StickerMessageContent
    for event in events:
        if isinstance(event, MessageEvent):
            if not isinstance(event.message, (TextMessageContent, ImageMessageContent, FileMessageContent, StickerMessageContent)):
                continue
            # 解析 session_id：依來源類型決定（user / group / room）
            source = event.source
            user_input = event.message.text if isinstance(event.message, TextMessageContent) else ""
            
            is_group_or_room = False
            if hasattr(source, "group_id") and source.group_id:
                session_id = f"line_group_{source.group_id}"
                chat_id = source.group_id
                is_group_or_room = True
            elif hasattr(source, "room_id") and source.room_id:
                session_id = f"line_room_{source.room_id}"
                chat_id = source.room_id
                is_group_or_room = True
            else:
                session_id = f"line_{source.user_id}"
                chat_id = source.user_id

            # Phase 1: Group Mention Filter & Window
            if is_group_or_room:
                if isinstance(event.message, TextMessageContent):
                    # 支援多種群組喚醒方式：[@Agent K], [@AgentK], @Agent K, @AgentK
                    mentions = ["[@Agent K]", "[@AgentK]", "@Agent K", "@AgentK"]
                    found_mention = False
                    
                    # 1. 首先進入快取 (不論是否叫它，都進快取供後續引用)
                    _add_to_cache(chat_id, event.message.id, user_input)
                    
                    for m in mentions:
                        if m in user_input:
                            user_input = user_input.replace(m, "").strip()
                            found_mention = True
                            break
                    
                    if found_mention:
                        import time
                        _last_request_time[f"mention_{chat_id}"] = time.time()
                    else:
                        # 不是叫它，直接忽略 (bypass processing)
                        logger.info(f"[LINE] Skipped group text (cached but no mention): chat={chat_id}")
                        continue
                else:
                    # For Image/File/Sticker in groups, check if bot was mentioned recently (window of 120s for better UX)
                    import time
                    last_mention = _last_request_time.get(f"mention_{chat_id}", 0)

                    # Phase 6: Proactive Cache
                    # If mentioned within 120s, we process it as a direct command
                    just_cache = (time.time() - last_mention > 120)
                    msg_type = type(event.message).__name__

                    if just_cache:
                        logger.info(f"[LINE] Group {msg_type} received without recent mention. Will only cache: chat={chat_id}")
                    else:
                        logger.info(f"[LINE] Group {msg_type} received with recent mention. Processing: chat={chat_id}")
                    
                    background_tasks.add_task(
                        _process_line_message,
                        line_api=line_api,
                        line_api_blob=line_api_blob,
                        reply_token=event.reply_token,
                        user_id=source.user_id,
                        chat_id=chat_id,
                        session_id=session_id,
                        event_msg=event.message,
                        extracted_text="",
                        quoted_file_path=None,
                        just_cache=just_cache
                    )
                    continue

            # Phase 6: Quote Recognition (引用識別)
            quoted_text = ""
            quoted_file = None
            if isinstance(event.message, TextMessageContent):
                # 嘗試取得引用訊息 ID
                quoted_msg_id = getattr(event.message, "quoted_message_id", None)
                if not quoted_msg_id:
                    m_dict = event.message.to_dict()
                    quoted_msg_id = m_dict.get("quotedMessageId")
                
                if quoted_msg_id:
                    q_data = _get_from_cache(chat_id, quoted_msg_id)
                    quoted_text = q_data.get("text")
                    quoted_file = q_data.get("file_path")
                    
                    if quoted_text:
                        logger.info(f"[LINE] Quoted text found in cache: {quoted_msg_id}")
                        user_input = f"[引用內容: \"{quoted_text}\"]\n{user_input}"
                    if quoted_file:
                        logger.info(f"[LINE] Quoted file found in cache: {quoted_file}")

            background_tasks.add_task(
                _process_line_message,
                line_api=line_api,
                line_api_blob=line_api_blob,
                reply_token=event.reply_token,
                user_id=source.user_id,
                chat_id=chat_id,
                session_id=session_id,
                event_msg=event.message,
                extracted_text=user_input,
                quoted_file_path=quoted_file
            )
            logger.info(
                f"[LINE] Queued background task: session={session_id}, "
                f"input='{user_input[:40]}...'"
            )

    # D. 立即回覆 200 OK — 不等待 LLM 完成
    return "OK"


# ── Proactive Messaging Endpoints ─────────────────────────────────────────────

@router.post("/api/line/push", tags=["Integration"])
async def line_push_message(request: Request):
    """
    主動推送訊息給特定 LINE 用戶。
    Payload: {"chat_id": "...", "text": "..."}
    """
    data = await request.json()
    chat_id = data.get("chat_id")
    text = data.get("text")

    if not chat_id or not text:
        raise HTTPException(status_code=400, detail="Missing chat_id or text")

    try:
        handler, line_api, line_api_blob = _get_line_components()
        _send_line_reply(line_api, None, chat_id, text)
        return {"status": "success", "message": f"Message pushed to {chat_id}"}
    except Exception as e:
        logger.error(f"[LINE] Push failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/line/broadcast", tags=["Integration"])
async def line_broadcast(request: Request):
    """
    向所有活躍的對話對象廣播文字訊息。
    Payload: {"text": "..."}
    """
    data = await request.json()
    text = data.get("text")
    if not text:
        raise HTTPException(status_code=400, detail="Missing text")

    try:
        handler, line_api, line_api_blob = _get_line_components()
        # 從快取的對話中收集 chat_id
        active_chats = set()
        for session_id in list(_last_request_time.keys()):
            if session_id.startswith("line_"):
                # line_U... -> U...
                chat_parts = session_id.split("_", 2)
                chat_id = chat_parts[-1]
                active_chats.add(chat_id)

        count = 0
        for chat_id in active_chats:
            try:
                _send_line_reply(line_api, None, chat_id, text)
                count += 1
            except:
                continue

        return {"status": "success", "broadcast_count": count}
    except Exception as e:
        logger.error(f"[LINE] Broadcast failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── Session Locking & UX ──────────────────────────────────────────────────────

_local_locks = {}
_local_lock_mutex = threading.Lock()
_last_request_time = {}  # 紀錄每個 session 的最後處理時間 (Debounce 用)

# ── Message Caching (Phase 6: Quote Support) ──────────────────────────────────
_message_cache = {}  # {chat_id: {msg_id: {"text": str, "file_path": str}}}
_MESSAGE_CACHE_LIMIT = 500  # 每回話上限 (LRU 簡化版)

def _add_to_cache(chat_id: str, msg_id: str, text: str = None, file_path: str = None):
    if chat_id not in _message_cache:
        _message_cache[chat_id] = {}
    cache = _message_cache[chat_id]
    
    # 如果已存在，更新非空欄位
    if msg_id in cache:
        if text: cache[msg_id]["text"] = text
        if file_path: cache[msg_id]["file_path"] = file_path
    else:
        cache[msg_id] = {"text": text, "file_path": file_path}
        
    if len(cache) > _MESSAGE_CACHE_LIMIT:
        oldest_key = next(iter(cache))
        del cache[oldest_key]

def _get_from_cache(chat_id: str, msg_id: str) -> dict:
    """回傳 {"text": str, "file_path": str}"""
    return _message_cache.get(chat_id, {}).get(msg_id, {"text": None, "file_path": None})

_redis_client = None

try:
    if redis is not None:
        _redis_client = redis.Redis(host="localhost", port=6379, db=0, socket_connect_timeout=1)
        _redis_client.ping()
        logger.info("[LINE] Redis connected for distributed locking.")
    else:
        logger.info("[LINE] Redis package not installed. Falling back to in-memory locks.")
except Exception as e:
    logger.info(f"[LINE] Redis not available ({e}). Falling back to in-memory locks.")
    _redis_client = None

@contextmanager
def _acquire_session_lock(session_id: str):
    """
    Acquire a lock for the session_id to prevent concurrent LLM requests.
    Uses Redis if available, else falls back to in-memory threading.Lock.
    """
    lock_key = f"lock:{session_id}"
    redis_lock = None
    local_lock = None

    if _redis_client:
        try:
            redis_lock = _redis_client.lock(lock_key, timeout=60)
            acquired = redis_lock.acquire(blocking=False)
            if not acquired:
                yield False
                return
        except Exception as e:
            logger.warning(f"[LINE] Redis lock failed, falling back to local: {e}")

    if not redis_lock:
        with _local_lock_mutex:
            if lock_key not in _local_locks:
                _local_locks[lock_key] = threading.Lock()
            local_lock = _local_locks[lock_key]
        
        acquired = local_lock.acquire(blocking=False)
        if not acquired:
            yield False
            return

    try:
        yield True
    finally:
        if redis_lock:
            try:
                redis_lock.release()
            except Exception:
                pass
        if local_lock:
            local_lock.release()

def _send_loading_animation(line_api, chat_id: str, seconds: int = 20):
    """呼叫 LINE Loading Animation API (使用官方 SDK)。可多次呼叫以延長動畫。"""
    from linebot.v3.messaging import ShowLoadingAnimationRequest

    try:
        # Note: LINE ShowLoadingAnimation only supports userId (Individual Chat).
        # It returns 400 Bad Request for groupId or roomId.
        if chat_id.startswith("U"):
            req = ShowLoadingAnimationRequest(chatId=chat_id, loadingSeconds=min(seconds, 60))
            line_api.show_loading_animation(req)
            logger.info(f"[LINE] Loading animation started for chat={chat_id} ({seconds}s)")
        else:
            logger.info(f"[LINE] Skipping loading animation for non-user chat={chat_id}")
    except Exception as e:
        logger.warning(f"[LINE] Exception starting loading animation: {e}")


def _send_status_push(line_api, chat_id: str, text: str):
    """推送中間狀態訊息給使用者（不佔用 reply_token）。"""
    try:
        from linebot.v3.messaging import TextMessage, PushMessageRequest
        line_api.push_message(
            PushMessageRequest(to=chat_id, messages=[TextMessage(text=text)])
        )
    except Exception as e:
        logger.warning(f"[LINE] Status push failed: {e}")


# ── Server-side File Content Extraction ───────────────────────────────────────

def _extract_file_content(file_path: str) -> tuple:
    """
    伺服器端預提取文件全文，讓 LLM 直接進行語意分析。
    支援多重 fallback：pdfplumber → pypdf、python-docx → docx2txt
    Returns: (full_text: str, error_message: str | None)
    """
    lower = file_path.lower()
    try:
        if lower.endswith('.docx'):
            text = _extract_docx(file_path)
        elif lower.endswith('.pdf'):
            text = _extract_pdf(file_path)
        elif lower.endswith(('.xlsx', '.xls')):
            import pandas as pd
            df = pd.read_excel(file_path)
            text = df.to_markdown(index=False)
        elif lower.endswith('.csv'):
            import pandas as pd
            df = pd.read_csv(file_path)
            text = df.to_markdown(index=False)
        else:
            # .txt, .md, .log, .json, .py, .js, .xml, etc.
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        return text.strip(), None
    except Exception as e:
        logger.error(f"[LINE] File content extraction failed: {file_path} → {e}")
        return "", str(e)


def _extract_pdf(file_path: str) -> str:
    """PDF 文字提取，依序嘗試 pdfplumber → pypdf"""
    # 優先使用 pdfplumber（表格 / 複雜版面效果最佳）
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if text.strip():
            logger.info("[LINE] PDF extracted via pdfplumber")
            return text
    except ImportError:
        logger.warning("[LINE] pdfplumber not installed, falling back to pypdf")
    except Exception as e:
        logger.warning(f"[LINE] pdfplumber failed ({e}), falling back to pypdf")

    # Fallback: pypdf（requirements.txt 中已列出）
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            logger.info("[LINE] PDF extracted via pypdf")
            return text
    except ImportError:
        logger.warning("[LINE] pypdf not installed either")
    except Exception as e:
        logger.warning(f"[LINE] pypdf also failed: {e}")

    raise RuntimeError("無法提取 PDF 內容：請確認伺服器已安裝 pdfplumber 或 pypdf")


def _extract_docx(file_path: str) -> str:
    """DOCX 文字提取，依序嘗試 python-docx → docx2txt"""
    # 優先使用 python-docx（結構化段落提取）
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 同時提取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        if text.strip():
            logger.info("[LINE] DOCX extracted via python-docx")
            return text
    except ImportError:
        logger.warning("[LINE] python-docx not installed, falling back to docx2txt")
    except Exception as e:
        logger.warning(f"[LINE] python-docx failed ({e}), falling back to docx2txt")

    # Fallback: docx2txt（requirements.txt 中已列出）
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            logger.info("[LINE] DOCX extracted via docx2txt")
            return text
    except ImportError:
        logger.warning("[LINE] docx2txt not installed either")
    except Exception as e:
        logger.warning(f"[LINE] docx2txt also failed: {e}")

    raise RuntimeError("無法提取 DOCX 內容：請確認伺服器已安裝 python-docx 或 docx2txt")


# ── Background Processing Function ────────────────────────────────────────────

def _preprocess_image(file_path: str) -> str:
    """實作 HEIC/RAW 自動轉 JPG 的預處理機制"""
    lower_path = file_path.lower()
    try:
        if lower_path.endswith(('.heic', '.heif')):
            import pillow_heif
            from PIL import Image
            heif_file = pillow_heif.read_heif(file_path)
            image = Image.frombytes(heif_file.mode, heif_file.size, heif_file.data)
            jpg_path = file_path.rsplit('.', 1)[0] + '.jpg'
            image.save(jpg_path, format="JPEG")
            return jpg_path
        elif lower_path.endswith(('.cr2', '.nef', '.arw', '.dng')):
            import rawpy
            from PIL import Image
            with rawpy.imread(file_path) as raw:
                rgb = raw.postprocess()
            image = Image.fromarray(rgb)
            jpg_path = file_path.rsplit('.', 1)[0] + '.jpg'
            image.save(jpg_path, format="JPEG")
            return jpg_path
    except Exception as e:
        logger.error(f"[LINE BG] Pre-process Failed for {file_path}: {e}")
    return file_path


def _download_sticker_image(sticker_id: str, uploads_dir: str, resource_type: str = "STATIC"):
    """
    嘗試從 LINE CDN 下載貼圖圖片，供 OpenAI Vision 辨識貼圖表情、動作與文字。
    針對不同貼圖類型 (STATIC / ANIMATION / POPUP) 使用不同的 CDN 路徑。
    回傳圖片路徑，失敗則回傳 None。
    """
    import os
    cached_path = os.path.join(uploads_dir, f"sticker_{sticker_id}.png")
    if os.path.exists(cached_path):
        return cached_path

    base = f"https://stickershop.line-scdn.net/stickershop/v1/sticker/{sticker_id}"

    # 依貼圖類型排列優先度：靜態 key frame 最清楚（文字可辨識），動態次之
    if resource_type in ("ANIMATION", "ANIMATION_SOUND", "POPUP", "POPUP_SOUND"):
        cdn_urls = [
            f"{base}/iPhone/sticker_key@2x.png",      # 動態貼圖的靜態 key frame（最清晰）
            f"{base}/iPhone/sticker@2x.png",           # 可能是 APNG，但 fallback 仍可用
            f"{base}/android/sticker.png",
        ]
    else:
        cdn_urls = [
            f"{base}/iPhone/sticker@2x.png",           # 標準靜態高解析度
            f"{base}/iPhone/sticker_key@2x.png",       # 部分靜態貼圖也有 key frame
            f"{base}/android/sticker.png",
            f"{base}/iPhone/sticker.png",
        ]

    for url in cdn_urls:
        try:
            resp = httpx.get(url, timeout=8, follow_redirects=True)
            if resp.status_code == 200 and len(resp.content) > 500:
                os.makedirs(uploads_dir, exist_ok=True)
                with open(cached_path, "wb") as f:
                    f.write(resp.content)
                logger.info(f"[LINE] Sticker image downloaded: id={sticker_id} type={resource_type} url={url} size={len(resp.content)}")
                return cached_path
        except Exception as e:
            logger.debug(f"[LINE] Sticker CDN failed ({url}): {e}")
            continue

    logger.warning(f"[LINE] All CDN URLs failed for sticker_id={sticker_id} type={resource_type}")
    return None


# ── Phase 8: Human-in-the-Loop — Pending State Handler ────────────────────────

def _handle_pending_state(
    line_api,
    chat_id: str,
    session_id: str,
    user_text: str,
) -> str | None:
    """
    檢查聊天是否有待確認的 pending state（approval 或 choice）。
    若使用者回覆的是確認/取消/選擇指令，處理 pending 操作並回傳回覆文字。
    若無 pending state 或使用者輸入不是確認指令，回傳 None（繼續正常流程）。

    Returns:
        str   → 已處理，回傳此文字給使用者
        None  → 無 pending 或非確認指令，繼續正常訊息處理
    """
    from server.core.pending_state import (
        PendingStateManager, parse_confirmation, parse_choice,
    )
    from main import PROJECT_ROOT

    pending_mgr = PendingStateManager(str(PROJECT_ROOT / "workspace" / "sessions"))
    pending = pending_mgr.get_pending(chat_id)

    if pending is None:
        return None  # 無待確認 → 正常流程

    pending_type = pending.get("type", "approval")

    # ══════════════════════════════════════════════════════════════════════
    # Type: "choice" — AI 提議了 A/B/C 方案，等使用者選擇
    # ══════════════════════════════════════════════════════════════════════
    if pending_type == "choice":
        valid_keys = [opt["key"] for opt in pending.get("options", [])]
        selected = parse_choice(user_text, valid_keys)

        if selected is None:
            # 使用者輸入不像選項 → 也檢查是否是取消
            confirmation = parse_confirmation(user_text)
            if confirmation == "reject":
                pending_mgr.clear_pending(chat_id)
                logger.info(f"[LINE HitL] User cancelled choice for chat={chat_id}")
                return "✅ 已取消方案選擇，有什麼需要可以繼續說！"
            # 不是選項也不是取消 → 清除 pending，視為新訊息
            pending_mgr.clear_pending(chat_id)
            logger.info(f"[LINE HitL] Choice pending cleared: unrelated input for chat={chat_id}")
            return None

        # 使用者選擇了某個方案
        chosen_option = next(
            (opt for opt in pending.get("options", []) if opt["key"] == selected),
            None,
        )
        chosen_text = chosen_option["text"] if chosen_option else selected
        original_query = pending.get("original_query", "")
        preamble = pending.get("preamble", "")

        logger.info(f"[LINE HitL] User selected choice={selected} ({chosen_text}) for chat={chat_id}")
        pending_mgr.clear_pending(chat_id)

        # 觸發 loading 動畫
        _send_loading_animation(line_api, chat_id, 60)

        try:
            from server.dependencies.uma import get_uma_instance
            from server.adapters.openai_adapter import OpenAIAdapter
            from server.dependencies.session import get_session_manager
            from server.services.runtime import get_universal_system_prompt

            _session_mgr = get_session_manager()
            _session_mgr.get_or_create_conversation(session_id, get_universal_system_prompt(platform="line"))
            _session_mgr._update_system_prompt(session_id, get_universal_system_prompt(platform="line"))

            uma = get_uma_instance()
            adapter = OpenAIAdapter(uma=uma)
            if not adapter.is_available:
                return f"⚠️ AI 服務暫時無法使用。你選擇了方案 {selected}：{chosen_text}"

            # 解析選項文字中的格式關鍵字，生成明確的格式指令
            _format_hint_map = {
                "pdf": "PDF (.pdf)",
                "word": "Word DOCX (.docx)",
                "docx": "Word DOCX (.docx)",
                "markdown": "Markdown (.md)",
                "md": "Markdown (.md)",
                "txt": "純文字 (.txt)",
                "excel": "Excel (.xlsx)",
                "xlsx": "Excel (.xlsx)",
                "csv": "CSV (.csv)",
                "pptx": "PowerPoint (.pptx)",
            }
            format_hint = ""
            chosen_lower = chosen_text.lower()
            for fmt_key, fmt_label in _format_hint_map.items():
                if fmt_key in chosen_lower:
                    format_hint = f"\n輸出檔案格式：{fmt_label}"
                    break

            # 注入使用者選擇作為 context，讓 LLM 根據選擇執行
            choice_msg = (
                f"[系統通知：使用者已從提議的方案中選擇了「{selected}. {chosen_text}」。\n"
                f"原始需求：{original_query}{format_hint}\n\n"
                f"請立即按照方案 {selected} 執行，使用 mcp-python-executor 工具建立檔案。"
                f"不要再次詢問或提議其他方案，直接行動並回覆結果。]"
            )
            _session_mgr.append_message(session_id, "user", choice_msg)

            history = _session_mgr.get_or_create_conversation(session_id)
            sys_msgs = [m for m in history if m.get("role") == "system"]
            rec_msgs = [m for m in history if m.get("role") != "system"][-20:]

            result_gen = adapter.chat(
                messages=sys_msgs + rec_msgs,
                user_query=choice_msg,
                session_id=session_id,
                tools_enabled=True,
            )

            final_reply = _collect_generator(
                result_gen,
                line_api=line_api,
                chat_id=chat_id,
                session_id=session_id,
            )
            return final_reply

        except Exception as e:
            logger.error(f"[LINE HitL] Choice execution failed: {e}", exc_info=True)
            return f"❌ 方案 {selected} 執行失敗：{e}"

    # ══════════════════════════════════════════════════════════════════════
    # Type: "approval" — 工具需要確認才能執行
    # ══════════════════════════════════════════════════════════════════════
    # 解析使用者回覆
    confirmation = parse_confirmation(user_text)

    if confirmation is None:
        # 使用者輸入不像確認/取消 → 清除過期 state，進入正常流程
        # 這讓使用者可以「忽略」待確認而繼續聊天
        pending_mgr.clear_pending(chat_id)
        logger.info(f"[LINE HitL] Pending cleared: user sent unrelated message for chat={chat_id}")
        return None

    # ── 使用者選擇「取消」 ──
    if confirmation == "reject":
        tool_name = pending.get("tool_name", "未知工具")
        pending_mgr.clear_pending(chat_id)
        logger.info(f"[LINE HitL] User rejected pending tool={tool_name} for chat={chat_id}")
        return f"✅ 已取消 `{tool_name}` 的執行。有其他需要可以繼續告訴我！"

    # ── 使用者選擇「確認」→ 執行待確認的工具 ──
    tool_name = pending.get("tool_name", "")
    tool_args = pending.get("tool_args", {})
    logger.info(f"[LINE HitL] User approved pending tool={tool_name} for chat={chat_id}")

    # 清除 pending（無論執行成功與否都不應重複執行）
    pending_mgr.clear_pending(chat_id)

    # 觸發 loading 動畫
    _send_loading_animation(line_api, chat_id, 60)

    try:
        from server.dependencies.uma import get_uma_instance
        from server.adapters.openai_adapter import OpenAIAdapter
        from server.dependencies.session import get_session_manager
        from server.services.runtime import get_universal_system_prompt
        import json

        uma = get_uma_instance()

        # 1. 執行被暫停的工具
        logger.info(f"[LINE HitL] Executing approved tool: {tool_name}({tool_args})")
        result = uma.execute_tool_call(tool_name, tool_args)
        result_text = json.dumps(result, ensure_ascii=False)
        logger.info(f"[LINE HitL] Tool result: {result_text[:200]}...")

        # 2. 把工具結果餵給 LLM，讓 AI 用自然語言回覆
        _session_mgr = get_session_manager()
        _session_mgr.get_or_create_conversation(session_id, get_universal_system_prompt(platform="line"))
        _session_mgr._update_system_prompt(session_id, get_universal_system_prompt(platform="line"))

        adapter = OpenAIAdapter(uma=uma)
        if not adapter.is_available:
            return f"⚠️ 工具 `{tool_name}` 已執行，但 AI 服務暫時無法生成回覆。\n\n原始結果：\n{result_text[:500]}"

        # 注入工具結果作為 assistant context
        tool_result_msg = (
            f"[系統通知：使用者已確認執行工具 `{tool_name}`。以下是執行結果：]\n\n"
            f"{result_text}\n\n"
            f"[請根據工具結果，用自然語言回覆使用者。不要說「使用者確認了」之類的流程說明。]"
        )
        _session_mgr.append_message(session_id, "user", tool_result_msg)

        history = _session_mgr.get_or_create_conversation(session_id)
        sys_msgs = [m for m in history if m.get("role") == "system"]
        rec_msgs = [m for m in history if m.get("role") != "system"][-20:]

        result_gen = adapter.chat(
            messages=sys_msgs + rec_msgs,
            user_query=tool_result_msg,
            session_id=session_id,
            tools_enabled=True,  # Allow follow-up tool calls if needed
        )

        final_reply = _collect_generator(
            result_gen,
            line_api=line_api,
            chat_id=chat_id,
            session_id=session_id,
        )
        return final_reply

    except Exception as e:
        logger.error(f"[LINE HitL] Tool execution failed: {e}", exc_info=True)
        return f"❌ 工具 `{tool_name}` 執行失敗：{e}"


def _process_line_message(
    line_api,
    line_api_blob,
    reply_token: str,
    user_id: str,
    chat_id: str,
    session_id: str,
    event_msg,
    extracted_text: str = "",
    quoted_file_path: str = None,
    just_cache: bool = False
):
    """
    背景函數：LLM 生成 → Tool 執行 → 組裝回覆 → 送回 LINE。

    重用現有元件：
    - router._session_mgr  (SessionManager)
    - OpenAIAdapter.chat() (full Tool Calling + RAG)
    - MEMORY.md 持久化（append_message 自動觸發）
    """
    from server.dependencies.uma import get_uma_instance
    from server.adapters.openai_adapter import OpenAIAdapter
    from server.dependencies.session import get_session_manager
    _session_mgr = get_session_manager()
    import time

    logger.info(f"[LINE BG] Start processing: session={session_id}")

    # 0. 防連點機制 (Debounce)：過濾 2 秒內重複觸發的事件
    current_time = time.time()
    last_time = _last_request_time.get(session_id, 0)
    if current_time - last_time < 2.0:
        logger.warning(f"[LINE BG] Debounced input for session={session_id} (Too fast)")
        return
    _last_request_time[session_id] = current_time

    with _acquire_session_lock(session_id) as acquired:
        if not acquired:
            logger.warning(f"[LINE BG] Session {session_id} is locked. Ignoring concurrent input.")
            return

        # 0.5 顯示 loading 動畫 (安撫使用者等待焦慮)，必須傳入 chat_id
        #     使用 60 秒以涵蓋檔案下載 + Tool Calling 耗時
        _send_loading_animation(line_api, chat_id, 60)

        try:
            from linebot.v3.webhooks import TextMessageContent, ImageMessageContent, FileMessageContent, StickerMessageContent
            import os

            # ── Phase 8: Human-in-the-Loop — Pending State Check ──────────────
            # 在處理新訊息前，先檢查此聊天是否有待確認的 pending state。
            # 若使用者回覆的是「確認」或「取消」，直接處理 pending 操作。
            if isinstance(event_msg, TextMessageContent) and extracted_text:
                pending_reply = _handle_pending_state(
                    line_api, chat_id, session_id, extracted_text
                )
                if pending_reply is not None:
                    # Pending state was handled — send reply and return
                    if len(pending_reply) > _LINE_MAX_CHARS:
                        pending_reply = pending_reply[:_LINE_MAX_CHARS] + "\n\n…（回覆過長，已截斷）"
                    _session_mgr.append_message(session_id, "user", extracted_text)
                    _session_mgr.append_message(session_id, "assistant", pending_reply)
                    _send_line_reply(line_api, reply_token, chat_id, pending_reply)
                    return

            attached_file_path = quoted_file_path
            _chunked_data = None  # Will be set if file exceeds 15,000 chars
            if isinstance(event_msg, TextMessageContent):
                user_input = extracted_text

                # ── Phase 2: Smart Clarification Injection ────────────────
                # 偵測「建立檔案/報告」意圖 → 注入釐清指令讓 LLM 語意分析
                # 缺少的資訊（地區、範圍、格式等），LLM 會一次問完
                if _detect_file_creation_intent(extracted_text):
                    user_input = (
                        f"[系統指令：使用者想要產出檔案或報告。在動手之前，請先分析需求是否完整。\n"
                        f"需要確認的關鍵資訊：\n"
                        f"- 主題的具體細節（如「天氣報告」→ 哪個地區？哪個時間範圍？）\n"
                        f"- 想要什麼檔案格式\n"
                        f"- 內容範圍與深度（簡要摘要 vs 詳細分析）\n\n"
                        f"規則：\n"
                        f"1. 如果缺少重要資訊，用簡潔友善的口吻「一次問完」所有問題\n"
                        f"2. 如果資訊已經完整明確，直接開始執行，不要多問\n"
                        f"3. 問問題時不要使用任何工具，純文字回覆即可\n"
                        f"4. 問題格式範例：「想先確認幾個細節：1) 哪個地區？2) 什麼格式？3) 要多詳細？」\n"
                        f"5. 嚴禁使用 [CHOICES] 格式，用自然語言編號列表即可]\n\n"
                        f"{extracted_text}"
                    )
                    logger.info(
                        f"[LINE HitL] Smart clarification injected for chat={chat_id}, "
                        f"query='{extracted_text[:50]}'"
                    )

                # 如果有引用檔案，注入特別提示
                if attached_file_path:
                    fname = os.path.basename(attached_file_path)
                    user_input = f"[系統通知：使用者引用了先前上傳的檔案 {fname}。檔案絕對路徑：{attached_file_path}]\n" + user_input

            elif isinstance(event_msg, StickerMessageContent):
                # ── Phase 7: Sticker Emotion Recognition ─────────────────────
                sticker_id = event_msg.sticker_id
                package_id = event_msg.package_id
                sticker_keywords = event_msg.keywords or []
                sticker_text = event_msg.text or ""
                resource_type = event_msg.sticker_resource_type or "STATIC"

                logger.info(
                    f"[LINE BG] Sticker received: pkg={package_id} id={sticker_id} "
                    f"keywords={sticker_keywords} text='{sticker_text}' type={resource_type}"
                )

                # 快取貼圖情緒資訊供後續引用
                cache_text = f"[貼圖: {', '.join(sticker_keywords)}]" if sticker_keywords else "[貼圖]"
                if sticker_text:
                    cache_text += f" {sticker_text}"
                _add_to_cache(chat_id, event_msg.id, text=cache_text)

                if just_cache:
                    logger.info(f"[LINE BG] Just cached sticker: session={session_id}, sticker_id={sticker_id}")
                    return

                # 嘗試下載貼圖圖片，供 Vision 模型辨識表情、動作與文字
                uploads_dir = os.path.join(os.getcwd(), "Agent_workspace", "line_uploads")
                sticker_image = _download_sticker_image(sticker_id, uploads_dir, resource_type)
                if sticker_image:
                    attached_file_path = sticker_image

                # 組裝情緒描述
                emotion_parts = []
                if sticker_keywords:
                    emotion_parts.append(f"情緒關鍵字：{', '.join(sticker_keywords)}")
                if sticker_text:
                    emotion_parts.append(f"貼圖中使用者自訂文字：「{sticker_text}」")
                emotion_desc = "；".join(emotion_parts) if emotion_parts else "（無 API 提供的情緒標籤）"

                has_image = "有" if sticker_image else "無"
                user_input = (
                    f"[系統通知：使用者傳送了一個 LINE 貼圖。\n"
                    f"貼圖後設資料：packageId={package_id}, stickerId={sticker_id}, 類型={resource_type}\n"
                    f"{emotion_desc}\n"
                    f"貼圖圖片：{has_image}\n\n"
                    f"【最重要 — 貼圖圖片 OCR】\n"
                    f"許多 LINE 貼圖的圖片上會印有中文、日文或英文字句（例如「辛苦了」「我罩你」「加油」等）。\n"
                    f"你必須仔細辨識圖片中所有文字並在回覆中自然融入，這是使用者表達的重點。\n"
                    f"如果圖片中有文字，請優先根據該文字的語意來理解使用者想表達什麼。\n\n"
                    f"【回覆風格】\n"
                    f"1. 觀察貼圖角色的表情、動作、姿態，加上圖中文字，綜合理解使用者情緒。\n"
                    f"2. 用活潑、口語化的語氣回覆，像朋友聊天一樣自然。\n"
                    f"3. 禁止說「你傳了一個貼圖」或複述系統通知內容。\n"
                    f"4. 禁止使用制式結語（如「如果有其他問題或需要協助的地方，隨時告訴我」之類的客套話）。\n"
                    f"5. 適當使用 emoji 讓回覆生動，但不過度。\n"
                    f"6. 結合對話上下文與貼圖情緒一起回覆。]"
                )

                _send_loading_animation(line_api, chat_id, 60)

            else:
                # Phase 2: Inbound Multimedia Download and Processing
                try:
                    import concurrent.futures
                    # Timeout protection: file download must complete within 30 seconds
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                        future = pool.submit(line_api_blob.get_message_content, event_msg.id)
                        try:
                            content_blob = future.result(timeout=30)
                        except concurrent.futures.TimeoutError:
                            raise TimeoutError(f"File download timed out after 30s (msg_id={event_msg.id})")

                    uploads_dir = os.path.join(os.getcwd(), "Agent_workspace", "line_uploads")
                    os.makedirs(uploads_dir, exist_ok=True)

                    if isinstance(event_msg, ImageMessageContent):
                        attached_file_path = os.path.join(uploads_dir, f"{event_msg.id}.jpg")
                        with open(attached_file_path, "wb") as f:
                            f.write(content_blob)

                        attached_file_path = _preprocess_image(attached_file_path)
                        user_input = "[系統通知：使用者傳送了一張圖片。請務必直接回答這張圖片裡面有什麼內容與細節，不要拒絕。]"

                    elif isinstance(event_msg, FileMessageContent):
                        filename = getattr(event_msg, "file_name", f"{event_msg.id}.bin")
                        attached_file_path = os.path.join(uploads_dir, f"{event_msg.id}_{filename}")
                        with open(attached_file_path, "wb") as f:
                            f.write(content_blob)

                        if attached_file_path.lower().endswith(('.heic', '.heif', '.cr2', '.nef', '.arw')):
                            attached_file_path = _preprocess_image(attached_file_path)
                            user_input = f"[系統通知：使用者傳送了一張高畫質圖片 {filename}，已轉為 {os.path.basename(attached_file_path)} 供您檢視]"
                        else:
                            # Server-side extraction: extract text content before sending to LLM
                            extracted_text, extract_err = _extract_file_content(attached_file_path)

                            if extract_err:
                                user_input = (
                                    f"[系統通知：使用者上傳了文件 {filename}，但伺服器無法提取內容。\n"
                                    f"錯誤訊息：{extract_err}\n"
                                    f"請告知使用者檔案可能已損壞、加密或格式不支援。]"
                                )
                            elif len(extracted_text) <= 15000:
                                # Single-pass mode: full content fits in one message
                                user_input = (
                                    f"[系統通知：使用者上傳了文件 {filename}。以下是完整內容：]\n\n"
                                    f"{extracted_text}\n\n"
                                    f"[請根據以上文件內容，直接進行分析、總結或處理使用者的需求。]"
                                )
                            else:
                                # Chunked mode: will be handled after session init
                                # Store chunks in a temporary variable for later processing
                                _chunked_data = {
                                    "filename": filename,
                                    "chunks": [extracted_text[i:i+15000] for i in range(0, len(extracted_text), 15000)]
                                }
                                user_input = None  # Will be set during chunked processing

                    # Re-trigger loading animation after download (timer may have expired)
                    _send_loading_animation(line_api, chat_id, 60)

                except Exception as e:
                    logger.error(f"[LINE BG] Download failed: {e}")
                    _send_status_push(line_api, chat_id, f"⚠️ 檔案下載失敗，請重新傳送。\n({e})")
                    return
                
                # 下載成功後補進快取，供後續引用
                if attached_file_path:
                    _add_to_cache(chat_id, event_msg.id, file_path=attached_file_path)

                if just_cache:
                    logger.info(f"[LINE BG] Just cached media: session={session_id}, msg_id={event_msg.id}")
                    return

            # 1. 取得或建立 Session（首次建立時注入 LINE 專屬 system prompt）
            # 注意：為了解決日期幻覺，偵測跨日對話並強制重置 OpenAI 狀態
            from datetime import datetime
            today_str = datetime.now().strftime("%Y-%m-%d")

            # Check for new day or new session in tracker
            if _session_days.get(session_id) != today_str:
                logger.info(f"[LINE BG] New day/session detected for {session_id} ({today_str}). Resetting OpenAI state.")
                _session_mgr.reset_openai_state(session_id)
                _session_days[session_id] = today_str

            _session_mgr.get_or_create_conversation(session_id, get_universal_system_prompt(platform="line"))
            _session_mgr._update_system_prompt(session_id, get_universal_system_prompt(platform="line"))

            # 2. 初始化 Adapter（LLM-as-a-Router + Token 節省配置）
            from server.services.model_router import (
                route_model, apply_token_fallback,
                get_model_chunk_summary, get_model_chunk_final,
            )
            uma = get_uma_instance()

            # Create a lightweight OpenAI client for the router (reuse credentials)
            _router_adapter = OpenAIAdapter(uma=uma)
            _openai_client = _router_adapter.client  # share the authenticated client

            # Route to optimal model tier based on task complexity
            _has_file = (attached_file_path is not None) or (_chunked_data is not None)
            _routed_model, _routed_tier = route_model(
                user_input=user_input or "",
                openai_client=_openai_client,
                has_file=_has_file,
                is_chunked_final=False,
            )

            adapter = OpenAIAdapter(uma=uma, model=_routed_model)
            # LINE 訊息上限 ~5000 字，OpenAI TPM 同時計算 max_output_tokens 預留量
            # 16384 → 2048 可節省 ~14336 tokens/request
            adapter.max_output_tokens = 2048
            logger.info(f"[LINE Router] Routed to model: {_routed_model} (tier={_routed_tier})")

            if not adapter.is_available:
                final_reply = "⚠️ AI 服務暫時無法使用，請確認 OPENAI_API_KEY 設定。"
            else:
                # ── Chunked Memory Processing ─────────────────────────────
                # For large files: process chunks sequentially, LLM summarizes
                # each chunk, then produces a final consolidated analysis.
                if user_input is None and _chunked_data is not None:
                    chunks = _chunked_data["chunks"]
                    fname = _chunked_data["filename"]
                    total = len(chunks)
                    logger.info(f"[LINE BG] Chunked mode: {fname} → {total} chunks")

                    # ── Per-stage model routing for chunks ──
                    # Intermediate summaries: use mini (cheap, 200K TPM, task is simple)
                    # Final synthesis: use full model (needs deep reasoning)
                    _chunk_summary_model = get_model_chunk_summary()
                    _chunk_final_model = get_model_chunk_final()
                    chunk_adapter = OpenAIAdapter(uma=uma, model=_chunk_summary_model)
                    chunk_adapter.max_output_tokens = 1024  # Summaries are short
                    logger.info(f"[LINE Router] Chunk summary model: {_chunk_summary_model}, final: {_chunk_final_model}")

                    # Accumulate ALL intermediate summaries in memory so the final
                    # synthesis call can receive the complete picture regardless of
                    # how aggressively the session-history token trimmer prunes old messages.
                    all_summaries: list[str] = []

                    for idx, chunk in enumerate(chunks):
                        part = idx + 1
                        if part < total:
                            # Intermediate chunk: ask LLM for concise summary
                            chunk_msg = (
                                f"[文件分段 {part}/{total}] 以下是 {fname} 的第 {part} 段內容：\n\n"
                                f"{chunk}\n\n"
                                f"[指令：閱讀並記住此段重點，以 500 字以內的摘要回覆。"
                                f"不要做最終分析，後續還有更多段落。]"
                            )
                            # Store only a short header in session — NOT the full 15000-char chunk
                            _session_mgr.append_message(session_id, "user", f"[文件分段 {part}/{total}：{fname}]")

                            history = _session_mgr.get_or_create_conversation(session_id)
                            sys_msgs = [m for m in history if m.get("role") == "system"]
                            # Pass only system prompt for chunk summarisation — no prior history
                            # needed and this keeps each chunk call lean and independent.
                            result_gen = chunk_adapter.chat(
                                messages=sys_msgs + [{"role": "user", "content": chunk_msg}],
                                user_query=chunk_msg,
                                session_id=session_id,
                                tools_enabled=False  # No tool calls for intermediate chunks
                            )
                            summary = _collect_generator(result_gen)
                            # Keep in session for conversational continuity
                            _session_mgr.append_message(session_id, "assistant", summary)
                            # Also keep in local accumulator — unaffected by token trimming
                            all_summaries.append(f"【第 {part} 段摘要】\n{summary}")
                            logger.info(f"[LINE BG] Chunk {part}/{total} summarized ({len(summary)} chars)")
                            _send_loading_animation(line_api, chat_id, 60)
                        else:
                            # Final chunk: switch to full model for deep synthesis
                            adapter = OpenAIAdapter(uma=uma, model=_chunk_final_model)
                            adapter.max_output_tokens = 2048
                            logger.info(f"[LINE Router] Switching to final model: {_chunk_final_model}")

                            # Build synthesis context from ALL accumulated summaries
                            # so the LLM receives the complete picture even if session history
                            # was trimmed by the token-aware budget below.
                            warning = ""
                            if total > 6:
                                warning = "\n⚠️ 文件過長，分析基於分段摘要，細節可能有遺漏。"

                            if all_summaries:
                                # Multi-chunk file: inject every summary explicitly
                                prior_context = "\n\n".join(all_summaries)
                                user_input = (
                                    f"以下是《{fname}》各段的摘要（共 {total} 段，最後一段的原文附於摘要之後）：\n\n"
                                    f"{prior_context}\n\n"
                                    f"【第 {part} 段（最終段）原文】\n{chunk}\n\n"
                                    f"[指令：根據以上所有段落摘要與最終段原文，"
                                    f"產出一份完整、深度的分析報告。{warning}]"
                                )
                            else:
                                # Single-chunk fallback (total == 1)
                                user_input = (
                                    f"[文件分段 {part}/{total} - 最終段] 以下是 {fname} 的最後一段內容：\n\n"
                                    f"{chunk}\n\n"
                                    f"[指令：這是文件的最後一段。請結合你先前記住的所有段落摘要，"
                                    f"產出一份完整、深度的分析報告。{warning}]"
                                )

                    # ── Clean up intermediate chunk entries from session ──────
                    # Remove the [文件分段 N/M：file] headers and their paired
                    # assistant summaries so they don't pollute future conversations.
                    # The final synthesis prompt already contains ALL summaries
                    # explicitly, so session no longer needs them.
                    _session_mgr.remove_chunk_entries(session_id, fname)

                    # Mark as chunked final → tools disabled for synthesis
                    _routed_tier = "chunk_final"

                # ── Normal Processing ─────────────────────────────────────
                # Safety: if user_input is still None after chunked processing, abort gracefully
                if user_input is None:
                    logger.error(f"[LINE BG] user_input is None after processing for session={session_id}")
                    final_reply = "⚠️ 檔案處理異常，請重新上傳。"
                    _send_line_reply(line_api, reply_token, chat_id, final_reply)
                    return

                # Append user message — store a TOKEN-SAFE version to prevent session bloat.
                # File content messages can be 10000+ chars; we only store a short header
                # so that subsequent requests don't carry gigabytes of history.
                _SESSION_MSG_CHAR_LIMIT = 400
                session_user_input = (
                    user_input[:_SESSION_MSG_CHAR_LIMIT] + "…[已截斷，完整內容僅用於本輪分析]"
                    if len(user_input) > _SESSION_MSG_CHAR_LIMIT else user_input
                )
                _session_mgr.append_message(session_id, "user", session_user_input)
                actual_input, execute_mode = _parse_command_prefix(user_input)

                history = _session_mgr.get_or_create_conversation(session_id)
                system_msgs = [m for m in history if m.get("role") == "system"]
                non_system = [m for m in history if m.get("role") != "system"]

                # Token-aware history trimming: accumulate from newest until budget is full.
                # Budget = 4000 tokens ≈ 12000 chars (leaving room for tools + output).
                _MAX_HISTORY_CHARS = 12000
                _acc_chars = 0
                trimmed_msgs = []
                for m in reversed(non_system):
                    mc = len(m.get("content", ""))
                    if _acc_chars + mc > _MAX_HISTORY_CHARS:
                        break
                    trimmed_msgs.insert(0, m)
                    _acc_chars += mc
                truncated_history = system_msgs + trimmed_msgs
                logger.info(f"[LINE BG] History: {len(non_system)} msgs → {len(trimmed_msgs)} msgs ({_acc_chars} chars)")

                # ── Token-based Fallback ────────────────────────────────
                # Pre-flight check: if estimated tokens exceed the current
                # model's safe budget, auto-downgrade to a cheaper model.
                _current_model = adapter.model
                _safe_model = apply_token_fallback(
                    model=_current_model,
                    messages=truncated_history,
                    max_output_tokens=adapter.max_output_tokens,
                )
                if _safe_model != _current_model:
                    adapter = OpenAIAdapter(uma=uma, model=_safe_model)
                    adapter.max_output_tokens = 2048
                    logger.info(f"[LINE Fallback] Downgraded: {_current_model} → {_safe_model}")

                # ── Tier-aware tool policy ─────────────────────────────
                # P1: Chunked final synthesis already has all content in prompt
                #     → no tool calls needed (saves a wasted API round-trip)
                # P2: nano tier = casual chat → tools off (saves ~800 tokens)
                #     mini tier = standard → max_tools=1 (light tool use)
                #     full/file tier = complex → full tool access
                if _routed_tier in ("chunk_final", "nano"):
                    _tools_enabled = False
                    _max_tools = 0
                elif _routed_tier == "mini":
                    _tools_enabled = execute_mode
                    _max_tools = 1
                else:  # full, file
                    _tools_enabled = execute_mode
                    _max_tools = 3

                result_gen = adapter.chat(
                    messages=truncated_history,
                    user_query=actual_input,
                    session_id=session_id,
                    attached_file=attached_file_path,
                    tools_enabled=_tools_enabled,
                    max_tools=_max_tools,
                )

                # Consume generator, assemble final reply
                final_reply = _collect_generator(
                    result_gen,
                    line_api=line_api,
                    chat_id=chat_id,
                    session_id=session_id,
                    user_input=user_input,
                )

            # 6. 截斷至 LINE 字元上限
            if len(final_reply) > _LINE_MAX_CHARS:
                final_reply = final_reply[:_LINE_MAX_CHARS] + "\n\n…（回覆過長，已截斷）"

            # 7. 寫入 Session 記憶（觸發 MEMORY.md 持久化）
            _session_mgr.append_message(session_id, "assistant", final_reply)

            # 8. 回傳 LINE（reply_token 優先，逾時後備 push_message）
            _send_line_reply(line_api, reply_token, chat_id, final_reply)

        except Exception as e:
            logger.error(
                f"[LINE BG] Unhandled error for session={session_id}: {e}", exc_info=True
            )
            _send_error_push(line_api, chat_id)


def _parse_command_prefix(user_input: str) -> tuple[str, bool]:
    """
    解析 LINE 訊息前綴指令，動態決定執行模式。

    /tool <msg>  → Agent 模式（強制 Tool Calling）
    /chat <msg>  → 純對話模式
    其他         → 預設 Agent 模式
    """
    if user_input.startswith("/tool "):
        return user_input[6:].strip(), True
    elif user_input.startswith("/chat "):
        return user_input[6:].strip(), False
    return user_input, True  # 預設啟用 Tool Calling


# ── Phase 2: Smart Clarification Detection ──────────────────────────────────

# File/report creation intent keywords
_FILE_INTENT_KW = [
    "整理成文件", "做成文件", "產出文件", "製作文件",
    "整理成檔案", "做成檔案", "產出檔案", "製作檔案",
    "做一份報告", "做份報告", "產出報告", "製作報告", "生成報告",
    "做一份文件", "建立文件", "建立檔案", "建立報告",
    "寫一份報告", "寫一份文件",
    "幫我做一份", "幫我製作", "幫我產出", "幫我建立",
]

# If user's request is already very detailed (has both format + topic details),
# skip the clarification injection and let LLM execute directly.
_COMPLETENESS_SIGNALS = [
    "pdf", "docx", "doc", "word", "xlsx", "excel", "csv",
    "txt", "pptx", "ppt", "md", "markdown", "json", "html",
]


def _detect_file_creation_intent(user_text: str) -> bool:
    """
    偵測使用者是否有建立檔案/報告的意圖，且需求可能不夠完整。

    Returns True → 注入釐清指令讓 LLM 先語意分析再決定問或做
    Returns False → 直接正常流程
    """
    # Must have file-creation intent
    has_intent = any(kw in user_text for kw in _FILE_INTENT_KW)
    if not has_intent:
        return False

    # If user already specified output format AND the text is long enough
    # (likely detailed), skip clarification to avoid unnecessary questions
    text_lower = user_text.lower()
    has_format = any(kw in text_lower for kw in _COMPLETENESS_SIGNALS)
    is_detailed = len(user_text) > 30  # Rough heuristic: short = likely vague

    if has_format and is_detailed:
        return False

    return True


def _collect_generator(
    result_gen,
    line_api=None,
    chat_id: str = None,
    session_id: str = None,
    user_input: str = None,
) -> str:
    """
    消費 adapter.chat() 的同步 generator，組裝完整回覆文字。

    當偵測到 Tool Call 中間狀態時，自動刷新 LINE loading 動畫，
    確保使用者在多輪 Tool Calling 期間持續看到 "..." 動畫。

    Human-in-the-Loop:
    - requires_approval → 儲存 approval pending state
    - [CHOICES]...[/CHOICES] → 儲存 choice pending state，讓使用者選方案

    Generator 的 chunk 格式：
    - {"status": "streaming", "content": "<partial text>"}
    - {"status": "success",   "content": "<full text>"}
    - {"status": "error",     "message": "<error msg>"}
    - {"status": "requires_approval", "tool_name": ..., "pending_args": ...}
    """
    accumulated = ""
    for chunk in result_gen:
        status = chunk.get("status")
        if status == "streaming":
            content = chunk.get("content", "")
            accumulated += content
            # Detect tool call marker from OpenAI adapter: "⚙️ 執行技能: ..."
            # Re-trigger loading animation so user keeps seeing "..." during long tool calls
            if line_api and chat_id and "⚙️" in content:
                _send_loading_animation(line_api, chat_id, 60)
        elif status == "success":
            # success chunk 包含完整最終內容
            final = chunk.get("content", "") or accumulated

            # ── Phase 2: Choice Detection ──────────────────────────
            # 檢查 LLM 回覆是否包含 [CHOICES]...[/CHOICES] 結構
            final = _maybe_save_choice_pending(
                final, chat_id=chat_id, session_id=session_id, user_input=user_input,
            )
            return final

        elif status == "error":
            err_msg = chunk.get("message", "未知錯誤")
            logger.error(f"[LINE BG] Adapter error: {err_msg}")
            return f"❌ 發生錯誤：{err_msg}"
        elif status == "requires_approval":
            tool_name = chunk.get("tool_name", "未知工具")
            pending_args = chunk.get("pending_args", {})
            risk_desc = chunk.get("risk_description", "高風險操作")

            # ── Human-in-the-Loop: 儲存 pending state ──────────────
            if chat_id:
                try:
                    from server.core.pending_state import PendingStateManager
                    from main import PROJECT_ROOT
                    pending_mgr = PendingStateManager(str(PROJECT_ROOT / "workspace" / "sessions"))
                    pending_mgr.set_pending(
                        chat_id=chat_id,
                        pending_type="approval",
                        data={
                            "tool_name": tool_name,
                            "tool_args": pending_args,
                            "risk_description": risk_desc,
                            "session_id": session_id or "",
                            "user_input": user_input or "",
                        },
                    )
                except Exception as e:
                    logger.error(f"[LINE BG] Failed to save pending state: {e}")

            return (
                f"⚠️ 工具 `{tool_name}` 需要你的確認才能執行。\n\n"
                f"📋 風險說明：{risk_desc}\n\n"
                f"請直接回覆：\n"
                f"  「確認」→ 執行此操作\n"
                f"  「取消」→ 放棄執行"
            )

    # streaming 結束但沒收到 success → 用 accumulated
    if accumulated:
        accumulated = _maybe_save_choice_pending(
            accumulated, chat_id=chat_id, session_id=session_id, user_input=user_input,
        )
        return accumulated

    return "（AI 未產生回覆，請稍後再試）"


def _maybe_save_choice_pending(
    llm_output: str,
    chat_id: str = None,
    session_id: str = None,
    user_input: str = None,
) -> str:
    """
    檢查 LLM 回覆是否包含 [CHOICES]...[/CHOICES] 區塊。
    若有，解析選項、儲存 choice pending state，並回傳 LINE 友好格式。
    若無，原文回傳。
    """
    if not chat_id or "[CHOICES]" not in llm_output:
        return llm_output

    try:
        from server.core.pending_state import (
            PendingStateManager, extract_choices, format_choices_for_line,
        )
        from main import PROJECT_ROOT

        parsed = extract_choices(llm_output)
        if parsed is None:
            # 格式不正確，直接回傳原文（不阻斷流程）
            return llm_output

        preamble, options = parsed
        logger.info(
            f"[LINE HitL] Choice proposal detected: {len(options)} options "
            f"for chat={chat_id}, keys={[o['key'] for o in options]}"
        )

        # 儲存 choice pending state
        pending_mgr = PendingStateManager(str(PROJECT_ROOT / "workspace" / "sessions"))
        pending_mgr.set_pending(
            chat_id=chat_id,
            pending_type="choice",
            data={
                "options": options,
                "preamble": preamble,
                "original_query": user_input or "",
                "session_id": session_id or "",
            },
        )

        # 回傳 LINE 友好格式（取代原始包含 [CHOICES] 標記的文字）
        return format_choices_for_line(preamble, options)

    except Exception as e:
        logger.error(f"[LINE HitL] Choice parsing failed: {e}", exc_info=True)
        return llm_output


def _send_line_reply(line_api, reply_token: str, chat_id: str, text: str):
    """
    發送回覆至 LINE。
    優先使用 reply_token（限 30 秒有效），逾時後備為 push_message。
    """
    from linebot.v3.messaging import TextMessage, ReplyMessageRequest, PushMessageRequest

    try:
        line_api.reply_message(
            ReplyMessageRequest(
                reply_token=reply_token,
                messages=[TextMessage(text=text)],
            )
        )
        logger.info(f"[LINE] Reply sent via reply_token → chat={chat_id}")
    except Exception as reply_err:
        # reply_token 已過期或失效，改用 push_message 主動推送
        logger.warning(
            f"[LINE] reply_token expired/failed ({reply_err}), "
            f"falling back to push_message → chat={chat_id}"
        )
        try:
            line_api.push_message(
                PushMessageRequest(
                    to=chat_id,
                    messages=[TextMessage(text=text)],
                )
            )
            logger.info(f"[LINE] Reply sent via push_message → chat={chat_id}")
        except Exception as push_err:
            logger.error(f"[LINE] push_message also failed: {push_err}")


def _send_error_push(line_api, chat_id: str):
    """發送通用錯誤通知至 LINE 使用者。"""
    try:
        from linebot.v3.messaging import TextMessage, PushMessageRequest

        line_api.push_message(
            PushMessageRequest(
                to=chat_id,
                messages=[
                    TextMessage(text="⚠️ 系統發生內部錯誤，請稍後再試或聯絡管理員。")
                ],
            )
        )
    except Exception as e:
        logger.error(f"[LINE] Failed to send error notification to {chat_id}: {e}")
