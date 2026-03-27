"""Shared file extraction helpers (text-first).

We reuse the same extraction policy used in LINE connector:
- PDF: pdfplumber -> pypdf
- DOCX: python-docx -> docx2txt
- XLSX/XLS/CSV: pandas to markdown
- Fallback: read as text with errors=replace

Images are intentionally not extracted here (by user request).
"""

from __future__ import annotations

from typing import Tuple


def extract_file_content(file_path: str) -> Tuple[str, str | None]:
    """Return (text, error)."""
    lower = file_path.lower()
    try:
        if lower.endswith(".docx"):
            text = _extract_docx(file_path)
        elif lower.endswith(".pdf"):
            text = _extract_pdf(file_path)
        elif lower.endswith((".xlsx", ".xls")):
            import pandas as pd

            df = pd.read_excel(file_path)
            text = df.to_markdown(index=False)
        elif lower.endswith(".csv"):
            import pandas as pd

            df = pd.read_csv(file_path)
            text = df.to_markdown(index=False)
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        return (text or "").strip(), None
    except Exception as e:
        return "", str(e)


def _extract_pdf(file_path: str) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if text.strip():
            return text
    except Exception:
        pass

    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        if text.strip():
            return text
    except Exception:
        pass

    import docx2txt

    return docx2txt.process(file_path) or ""
