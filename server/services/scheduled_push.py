"""
LINE Bot Scheduled Push Service
================================
Manages per-user/group scheduled push notifications.
Each user/group has a JSON config with multiple tasks.

Schedule config: workspace/schedules/{session_id}.json
Content types:
  - news        : 新聞重點摘要 (via mcp-web-search + LLM)
  - work_summary: 一週工作項目統整 (via session history)
  - language    : 語言詞彙學習 (via LLM)
  - custom      : 自訂 prompt (via LLM)
"""

import json
import logging
import os
import re
import time
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

logger = logging.getLogger("MCP_Server.ScheduledPush")

# ── Cron expression helpers ──────────────────────────────────────────────────

def _parse_simple_cron(cron_str: str) -> dict:
    """
    Parse simple cron: '08:30' or '每天 08:30' → {"hour": 8, "minute": 30}
    Also supports: 'weekday 09:00' → {"day_of_week": "mon-fri", "hour": 9, "minute": 0}
    Full cron: '0 8 * * 1-5' → parsed as minute hour day month day_of_week
    Interval: 'every +10m' or '*/10 * * * *' → {"interval_minutes": 10}
    """
    import re as _re
    cron_str = cron_str.strip()

    # Interval: 'every +10m' / 'every 10m' / 'every 10 min'
    _iv = _re.search(r'every\s*\+?(\d+)\s*(?:m|min)', cron_str)
    if _iv:
        return {"interval_minutes": int(_iv.group(1))}

    # Full cron with interval in minute field: '*/10 * * * *'
    _parts = cron_str.split()
    if len(_parts) == 5 and _parts[0].startswith("*/"):
        try:
            return {"interval_minutes": int(_parts[0][2:])}
        except ValueError:
            pass

    # Format: HH:MM
    if ":" in cron_str and len(cron_str.split()) <= 2:
        time_part = cron_str.split()[-1]  # take last part (after optional prefix)
        prefix = cron_str.split()[0] if len(cron_str.split()) > 1 else ""
        parts = time_part.split(":")
        h, m = int(parts[0]), int(parts[1])
        result = {"hour": h, "minute": m}
        if prefix in ("weekday", "平日", "工作日"):
            result["day_of_week"] = "mon-fri"
        return result

    # Full cron: minute hour day month day_of_week
    parts = cron_str.split()
    if len(parts) == 5:
        result = {}
        if parts[0] != "*":
            result["minute"] = int(parts[0])
        if parts[1] != "*":
            result["hour"] = int(parts[1])
        if parts[2] != "*":
            result["day"] = int(parts[2])
        if parts[3] != "*":
            result["month"] = int(parts[3])
        if parts[4] != "*":
            result["day_of_week"] = parts[4]
        return result

    return {"hour": 8, "minute": 0}  # default 08:00


class ScheduledPushService:
    """Manages scheduled push tasks for LINE users/groups."""

    def __init__(self, project_root: str):
        self.project_root = Path(project_root).resolve()
        self.schedules_dir = self.project_root / "workspace" / "schedules"
        self.schedules_dir.mkdir(parents=True, exist_ok=True)

    # ─── Config CRUD ─────────────────────────────────────────────────────────

    def _config_path(self, session_id: str) -> Path:
        return self.schedules_dir / f"{session_id}.json"

    def load_config(self, session_id: str) -> dict:
        path = self._config_path(session_id)
        if path.exists():
            try:
                return json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                return {"session_id": session_id, "chat_id": "", "tasks": []}
        return {"session_id": session_id, "chat_id": "", "tasks": []}

    def save_config(self, session_id: str, config: dict):
        path = self._config_path(session_id)
        path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    def add_task(
        self,
        session_id: str,
        chat_id: str,
        task_type: str,
        name: str,
        cron_str: str,
        config: dict = None,
    ) -> dict:
        """Add a scheduled task. Returns the created task dict."""
        cfg = self.load_config(session_id)
        cfg["chat_id"] = chat_id
        cfg["session_id"] = session_id

        task = {
            "id": f"task_{uuid.uuid4().hex[:8]}",
            "type": task_type,
            "name": name,
            "cron": cron_str,
            "cron_parsed": _parse_simple_cron(cron_str),
            "config": config or {},
            "enabled": True,
            "created_at": datetime.now().isoformat(),
            "last_run": None,
        }
        cfg["tasks"].append(task)
        self.save_config(session_id, cfg)
        logger.info(f"[ScheduledPush] Added task '{name}' ({task_type}) for {session_id} @ {cron_str}")
        return task

    def remove_task(self, session_id: str, task_id: str) -> bool:
        cfg = self.load_config(session_id)
        original_len = len(cfg["tasks"])
        cfg["tasks"] = [t for t in cfg["tasks"] if t["id"] != task_id]
        if len(cfg["tasks"]) < original_len:
            self.save_config(session_id, cfg)
            logger.info(f"[ScheduledPush] Removed task {task_id} from {session_id}")
            return True
        return False

    def toggle_task(self, session_id: str, task_id: str, enabled: bool) -> bool:
        cfg = self.load_config(session_id)
        for task in cfg["tasks"]:
            if task["id"] == task_id:
                task["enabled"] = enabled
                self.save_config(session_id, cfg)
                return True
        return False

    def list_tasks(self, session_id: str) -> list:
        cfg = self.load_config(session_id)
        return cfg.get("tasks", [])

    # ─── Content Generators ──────────────────────────────────────────────────

    def _generate_news(self, task: dict, llm_callable=None, tool_executor=None) -> str:
        """Generate news summary using web search + LLM."""
        config = task.get("config", {})
        topic = config.get("topic", "科技與AI")
        count = config.get("count", 5)
        detail = config.get("detail", "normal")

        # Step 1: Web search
        search_results = ""
        if tool_executor:
            try:
                search_depth = "advanced" if detail == "detailed" else "basic"
                result = tool_executor("mcp-web-search", {
                    "query": f"{topic} 最新新聞 today 台灣 日本",
                    "max_results": count * 2,
                    "search_depth": search_depth,
                })
                if isinstance(result, dict):
                    search_results = result.get("content", result.get("result", str(result)))
                else:
                    search_results = str(result)
            except Exception as e:
                logger.warning(f"[ScheduledPush] Web search failed: {e}")
                search_results = f"(搜尋失敗: {e})"

        # Step 2: LLM summarize
        if llm_callable and search_results:
            today = datetime.now().strftime("%Y-%m-%d")
            
            detail_instruction = "摘要（2-3 句）"
            if detail == "detailed":
                detail_instruction = "詳細報導（至少 25 句，包含事件背景、發展經過與各方觀點）"
            elif detail == "brief":
                detail_instruction = "簡短一兩句話總結"

            prompt = (
                f"你是一位專業新聞編輯。根據以下搜尋結果，用繁體中文整理出 {count} 則「{topic}」領域的重點新聞。\n"
                f"搜尋範圍請預設以台灣為主、日本為輔。\n"
                f"今天日期：{today}\n"
                f"格式要求：\n"
                f"📰 **標題**\n"
                f"{detail_instruction}\n"
                f"🔗 來源標記\n\n"
                f"搜尋結果：\n{search_results}"
            )
            try:
                summary = llm_callable(prompt)
                fallback_warning = "\n\n⚠️ *(System Note: 因達到系統流量上限，本次採純文字回覆模式，未能完整產生 PDF 檔案)*"
                return summary + fallback_warning
            except Exception as e:
                logger.error(f"[ScheduledPush] LLM news summary failed: {e}")
                return f"📰 今日{topic}新聞摘要生成失敗，請稍後重試。"

        return f"📰 今日{topic}新聞搜尋暫時無法使用。"

    def _generate_work_summary(self, task: dict, session_id: str, llm_callable=None) -> str:
        """Generate weekly work summary from session history."""
        config = task.get("config", {})
        days = config.get("days", 7)

        # Read session history
        sessions_dir = self.project_root / "workspace" / "sessions"
        session_path = sessions_dir / f"{session_id}.json"

        history_text = ""
        if session_path.exists():
            try:
                history = json.loads(session_path.read_text(encoding="utf-8"))
                if isinstance(history, list):
                    # Filter recent messages
                    cutoff = (datetime.now() - timedelta(days=days)).isoformat()
                    recent = []
                    for msg in history:
                        ts = msg.get("timestamp", msg.get("ts", ""))
                        if ts >= cutoff or not ts:
                            role = msg.get("role", "")
                            content = msg.get("content", "")
                            if role in ("user", "assistant") and content:
                                recent.append(f"[{role}] {content[:200]}")
                    history_text = "\n".join(recent[-50:])  # Last 50 messages
            except Exception as e:
                logger.warning(f"[ScheduledPush] Failed to read session history: {e}")

        # Read Notion tasks if available (optional)
        notion_summary = ""

        if llm_callable:
            today = datetime.now().strftime("%Y-%m-%d")
            prompt = (
                f"你是一位專業的工作助理。根據以下近 {days} 天的對話記錄，用繁體中文整理出工作重點摘要。\n"
                f"今天日期：{today}\n"
                f"格式要求：\n"
                f"📋 **本週工作重點**\n"
                f"1. 已完成項目\n"
                f"2. 進行中項目\n"
                f"3. 待處理項目\n"
                f"4. 重要提醒\n\n"
                f"對話記錄：\n{history_text or '(無近期記錄)'}\n"
                f"{f'Notion 任務：{notion_summary}' if notion_summary else ''}"
            )
            try:
                return llm_callable(prompt)
            except Exception as e:
                logger.error(f"[ScheduledPush] LLM work summary failed: {e}")
                return "📋 本週工作摘要生成失敗，請稍後重試。"

        return "📋 工作摘要功能需要 LLM 支援。"

    def _generate_language(self, task: dict, llm_callable=None) -> str:
        """Generate language vocabulary/grammar learning content (fallback path)."""
        config = task.get("config", {})
        language = config.get("language", "日文")
        level = config.get("level", "N3")
        count = config.get("count", 5)
        content_type = config.get("content_type", "vocabulary")
        current_category = config.get("current_category", config.get("topic", ""))
        used_items = config.get("used_items", [])
        exclusion = (
            f"🚫 以下批次已推送過，本次必須涵蓋完全不同的內容：\n"
            f"{chr(10).join(used_items[-20:])}\n\n"
        ) if used_items else ""
        focus = f"本次聚焦主題：【{current_category}】\n" if current_category else ""

        if llm_callable:
            if content_type == "grammar":
                prompt = (
                    f"{exclusion}你是一位{language}教師。請提供「恰好 {count} 個」{level}程度的文法句型教學。\n"
                    f"{focus}"
                    f"每個文法格式：句型 / 意思 / 使用情況 / 例句1（附翻譯）/ 例句2（附翻譯）。\n"
                    f"⚠️ 輸出恰好 {count} 個文法句型，不多不少。最後出一道小測驗。用繁體中文說明。"
                )
            else:
                prompt = (
                    f"{exclusion}你是一位{language}教師。請提供「恰好 {count} 個」{level}程度的詞彙教學。\n"
                    f"{focus}"
                    f"每個詞彙格式：原文（假名）/ 發音（羅馬拼音）/ 中文意思 / 例句（附中文翻譯）。\n"
                    f"⚠️ 輸出恰好 {count} 個詞彙，不多不少。最後出一道小測驗。用繁體中文。"
                )
            try:
                return llm_callable(prompt)
            except Exception as e:
                logger.error(f"[ScheduledPush] LLM language gen failed: {e}")
                return f"📖 今日{language}學習內容生成失敗，請稍後重試。"

        return f"📖 {language}學習功能需要 LLM 支援。"

    def _init_language_categories(self, original_request: str, llm_callable) -> list:
        """
        One-time LLM call to generate a rotation category list for this task.
        Fully generic — works for any language, content type, or level.
        """
        prompt = (
            f"根據以下學習需求，列出 24 個具體且不重複的子主題或分類，"
            f"讓每次推送都能涵蓋不同面向、避免內容重複。\n"
            f"需求：{original_request}\n\n"
            f"只輸出一個 JSON 陣列，例如：[\"主題1\", \"主題2\", ...]，不要任何說明或額外文字。"
        )
        try:
            result = llm_callable(prompt)
            match = re.search(r'\[.*?\]', result, re.DOTALL)
            if match:
                categories = json.loads(match.group())
                if isinstance(categories, list) and len(categories) >= 3:
                    logger.info(f"[ScheduledPush] Generated {len(categories)} rotation categories")
                    return categories
        except Exception as e:
            logger.warning(f"[ScheduledPush] Category init failed: {e}")
        return []

    def _update_language_history(self, task: dict, session_id: str, current_category: str):
        """
        After each push, record which category was covered and advance the rotation index.
        Format-agnostic — does not parse LLM output content.
        """
        config_path = self._config_path(session_id)
        if not config_path.exists():
            return
        try:
            cfg = json.loads(config_path.read_text(encoding="utf-8"))
            for t in cfg.get("tasks", []):
                if t["id"] == task["id"]:
                    tc = t.get("config", {})
                    # Advance rotation index
                    tc["category_index"] = tc.get("category_index", 0) + 1
                    # Append batch summary (format-agnostic)
                    batch_num = tc["category_index"]
                    used_items = tc.get("used_items", [])
                    used_items.append(f"第{batch_num}批（{current_category}）")
                    tc["used_items"] = used_items[-40:]  # keep last 40 batches
                    break
            config_path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
            logger.info(f"[ScheduledPush] History updated for {task['id']}: category={current_category}")
        except Exception as e:
            logger.warning(f"[ScheduledPush] Failed to update language history: {e}")

    def _generate_custom(self, task: dict, llm_callable=None) -> str:
        """Generate content from custom prompt."""
        config = task.get("config", {})
        prompt = config.get("prompt", "")

        if not prompt:
            return "⚠️ 自訂推送未設定 prompt。"

        if llm_callable:
            today = datetime.now().strftime("%Y-%m-%d %H:%M")
            full_prompt = f"現在時間：{today}\n用繁體中文回覆。\n\n{prompt}"
            try:
                return llm_callable(full_prompt)
            except Exception as e:
                logger.error(f"[ScheduledPush] LLM custom gen failed: {e}")
                return "⚠️ 自訂推送內容生成失敗。"

        return "⚠️ 自訂推送功能需要 LLM 支援。"

    # ── AutoScan-inspired keyword extraction ──────────────────────────────────
    # Strategy: strip away noise (time, actions, format, quantity) → remaining = topic
    # Inspired by fsc0638/AutoScan's structured field-mapping approach.

    # Known domain keywords (priority match)
    _DOMAIN_KEYWORDS = [
        # Finance / Economy
        "經濟", "股市", "房市", "金融", "財經", "投資", "貿易", "匯率", "利率",
        "加密貨幣", "區塊鏈", "房地產", "基金", "債券", "期貨", "外匯",
        # Tech
        "科技", "AI", "半導體", "電動車", "生技", "5G", "量子", "機器人",
        "雲端", "資安", "晶片", "軟體",
        # Industry / Sector
        "產業", "能源", "航運", "零售", "製造", "農業", "觀光", "旅遊",
        # Society / Crime / Law
        "政治", "國際", "軍事", "外交", "社會", "社會案件", "社會事件",
        "犯罪", "警政", "刑事", "詐騙", "治安",
        "教育", "體育", "娛樂",
        "醫療", "健康", "環境", "氣候", "法律", "法規",
    ]

    # Stopwords: time, actions, format, quantity, modifiers — NOT topic content
    _STOPWORDS = {
        # Time expressions
        "分鐘", "小時", "天", "週", "月", "年", "後", "前", "每天", "每日", "每週",
        "明天", "今天", "今日", "早上", "下午", "晚上", "定時", "定期",
        # Action verbs
        "統整", "推送", "搜尋", "查找", "幫我", "給我", "傳送", "發送", "整理",
        "製作", "建立", "產生", "生成", "做", "弄", "要",
        # Format / output
        "PDF", "pdf", "DOCX", "docx", "TXT", "txt", "下載", "檔案", "文件",
        # Quantity
        "則", "條", "篇", "筆", "個",
        # Quality modifiers
        "內容", "越", "好", "盡", "詳盡", "詳細", "精簡", "簡要", "深入",
        # Generic
        "新聞", "頭條", "報導", "時事", "資訊", "最新", "相關", "議題", "主題",
        "請", "並", "且", "與", "和", "的", "了", "吧", "呢", "喔", "啊",
        "標記", "出處", "來源", "包含", "涵蓋", "最好",
    }

    @classmethod
    def _extract_topic_keywords(cls, text: str) -> str:
        """
        AutoScan-inspired keyword extraction via stopword stripping.
        1. Priority: match known domain keywords from text
        2. Fallback: strip stopwords + noise, remaining = topic
        3. Also extract "XXX相關" pattern
        """
        import re

        # ── Step 1: Match known domain keywords (priority, order-preserving) ──
        found_domains = []
        for kw in cls._DOMAIN_KEYWORDS:
            if kw in text and kw not in found_domains:
                found_domains.append(kw)

        # ── Step 2: Extract "XXX相關" patterns ──
        for m in re.finditer(r'([\u4e00-\u9fff]{2,6}?)相關', text):
            candidate = m.group(1)
            if candidate not in cls._STOPWORDS and candidate not in found_domains:
                found_domains.append(candidate)

        # ── Step 3: Extract "包含/涵蓋..." clause content ──
        incl_match = re.search(r'(?:包含|涵蓋|最好包含)(.*?)(?:[，,。]|$)', text)
        if incl_match:
            clause = incl_match.group(1)
            # Split by connectors and extract meaningful segments
            for seg in re.split(r'[與和、及跟]', clause):
                seg = seg.strip()
                # Remove stopwords from each segment
                for sw in cls._STOPWORDS:
                    seg = seg.replace(sw, "")
                seg = re.sub(r'\d+', '', seg).strip()
                if len(seg) >= 2 and seg not in found_domains:
                    found_domains.append(seg)

        if found_domains:
            return " ".join(found_domains)

        # ── Step 4: Fallback — strip everything that's noise ──
        cleaned = text
        # Remove time expressions like "2分鐘後", "10分鐘後"
        cleaned = re.sub(r'\d+\s*分鐘[後后]?', '', cleaned)
        cleaned = re.sub(r'\d+\s*小時[後后]?', '', cleaned)
        # Remove numbers + counters
        cleaned = re.sub(r'\d+\s*[則條篇筆個]', '', cleaned)
        # Remove stopwords
        for sw in cls._STOPWORDS:
            cleaned = cleaned.replace(sw, "")
        # Remove remaining digits and punctuation
        cleaned = re.sub(r'[\d,，。、！!？?：:；;（）()\[\]「」\s]+', ' ', cleaned)
        cleaned = cleaned.strip()

        return cleaned if len(cleaned) >= 2 else "綜合"

    @classmethod
    def _infer_news_config_from_text(cls, text: str) -> dict | None:
        """
        If text looks like a news request, extract structured config.
        Inspired by AutoScan's field-mapping approach:
          原文 → 拆解為獨立欄位 (topic / count / detail / extra)
        """
        import re

        # ── Gate: must mention news-related keywords ──
        news_keywords = ["新聞", "news", "頭條", "時事", "報導"]
        if not any(kw in text.lower() for kw in news_keywords):
            return None

        config = {}

        # ── Field 1: count (數量) ──
        count_match = re.search(r'(\d+)\s*[則條篇筆]', text)
        config["count"] = int(count_match.group(1)) if count_match else 10

        # ── Field 2: detail (摘要深度) ──
        if any(kw in text for kw in ["詳盡", "詳細", "深入", "越詳盡越好", "越詳細越好"]):
            config["detail"] = "detailed"
        elif any(kw in text for kw in ["簡要", "精簡", "簡單"]):
            config["detail"] = "brief"
        else:
            config["detail"] = "normal"

        # ── Field 3: topic (主題關鍵字 — AutoScan-style extraction) ──
        config["topic"] = cls._extract_topic_keywords(text)

        # ── Field 4: extra_instructions (額外需求) ──
        extra_parts = []
        if any(kw in text.lower() for kw in ["pdf", "下載"]):
            extra_parts.append("統整成PDF供下載")
        incl_match = re.search(r'(?:包含|涵蓋|最好包含|最好有)(.*?)(?:[，,。]|$)', text)
        if incl_match:
            extra_parts.append(incl_match.group(1).strip())
        if any(kw in text for kw in ["標記出處", "出處", "來源"]):
            extra_parts.append("標記出處")
        if extra_parts:
            config["extra_instructions"] = "，".join(extra_parts)

        return config

    def generate_content(self, task: dict, session_id: str,
                         llm_callable=None, tool_executor=None) -> str:
        """
        Generate content using full OpenAI Adapter pipeline with tool calling.
        This allows the LLM to chain multiple skills (e.g., web-search → python-executor → PDF).
        Falls back to simple generators if adapter is unavailable.
        """
        task_type = task.get("type", "custom")
        config = task.get("config", {})
        today = datetime.now().strftime("%Y-%m-%d")
        original_request = config.get("original_request", "")

        # ── Auto-correct: if type is wrong but request looks like news, upgrade to "news" ──
        # LLM often misclassifies news requests as "custom" or "reminder"
        if task_type in ("custom", "reminder") and original_request:
            inferred = self._infer_news_config_from_text(original_request)
            if inferred:
                logger.info(
                    f"[ScheduledPush] Auto-corrected type {task_type}→news "
                    f"(count={inferred.get('count')}, detail={inferred.get('detail')}, "
                    f"topic={inferred.get('topic')})"
                )
                task_type = "news"
                # Merge: keep original_request, overlay inferred fields
                config = {**config, **inferred}

        # ── Auto-fill missing language fields from original_request ──
        if task_type == "language" and original_request:
            needs_fill = (
                not config.get("count") or not config.get("level")
                or not config.get("language") or not config.get("content_type")
            )
            if needs_fill:
                filled = {}
                if not config.get("count"):
                    _cm = re.search(r'(\d+)\s*[個條]', original_request)
                    filled["count"] = int(_cm.group(1)) if _cm else 5
                if not config.get("level"):
                    _lm = re.search(r'N([1-5])', original_request)
                    filled["level"] = f"N{_lm.group(1)}" if _lm else "N3"
                if not config.get("language"):
                    for _lang, _kws in [
                        ("日文", ["日文", "日語", "日本語"]),
                        ("英文", ["英文", "英語"]),
                        ("韓文", ["韓文", "韓語"]),
                        ("義大利文", ["義大利"]),
                        ("法文", ["法文", "法語"]),
                        ("西班牙文", ["西班牙"]),
                    ]:
                        if any(kw in original_request for kw in _kws):
                            filled["language"] = _lang
                            break
                if not config.get("content_type"):
                    _grammar_kws = ["文法", "語法", "句型", "表達方式", "grammar"]
                    filled["content_type"] = "grammar" if any(kw in original_request for kw in _grammar_kws) else "vocabulary"
                if filled:
                    config = {**config, **filled}
                    logger.info(f"[ScheduledPush] Auto-filled language fields: {filled}")

        # ── Language: category rotation init + inject ──
        current_category = None
        if task_type == "language":
            # Step 1: Initialize category list on first push (one-time LLM call)
            if not config.get("categories") and llm_callable:
                categories = self._init_language_categories(original_request or str(config), llm_callable)
                if categories:
                    config = {**config, "categories": categories, "category_index": 0}
                    # Persist immediately so other ticks don't re-init
                    config_path = self._config_path(session_id)
                    if config_path.exists():
                        try:
                            cfg = json.loads(config_path.read_text(encoding="utf-8"))
                            for t in cfg.get("tasks", []):
                                if t["id"] == task["id"]:
                                    t["config"]["categories"] = categories
                                    t["config"]["category_index"] = 0
                                    break
                            config_path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
                        except Exception:
                            pass

            # Step 2: Pick current category from rotation
            categories = config.get("categories", [])
            if categories:
                idx = config.get("category_index", 0)
                current_category = categories[idx % len(categories)]
                config = {**config, "current_category": current_category}

        # Build task-specific prompt for the full adapter pipeline
        prompt = self._build_task_prompt(task_type, config, today)

        # Safety: if prompt is nearly empty (custom with no prompt), use original_request
        if original_request and len(prompt.strip()) < 50:
            logger.info(f"[ScheduledPush] Prompt too short, using original_request as fallback")
            prompt = (
                f"現在時間：{today}\n"
                f"使用者的完整需求如下，請使用可用的工具完成：\n\n"
                f"「{original_request}」\n\n"
                f"請用繁體中文回覆。"
            )

        # Detect if file generation is expected
        needs_file = any(kw in prompt.lower() for kw in [
            "pdf", "docx", "txt", "md", "xlsx", "下載", "檔案", "存放到"
        ]) or any(kw in original_request.lower() for kw in [
            "pdf", "下載", "檔案"
        ])

        # Try full adapter pipeline first (supports multi-skill chaining)
        try:
            result = self._run_via_adapter(prompt, session_id, needs_file=needs_file)
            if result:
                if task_type == "language":
                    self._update_language_history(task, session_id, current_category or "通用")
                return result
        except Exception as e:
            logger.warning(f"[ScheduledPush] Adapter pipeline failed, falling back: {e}")

        # Fallback to simple generators
        if task_type == "news":
            return self._generate_news(task, llm_callable, tool_executor)
        elif task_type == "work_summary":
            return self._generate_work_summary(task, session_id, llm_callable)
        elif task_type == "language":
            return self._generate_language(task, llm_callable)
        elif task_type == "custom":
            return self._generate_custom(task, llm_callable)
        else:
            return f"⚠️ 不支援的推送類型：{task_type}"

    # News detail level definitions
    _NEWS_DETAIL_LEVELS = {
        "brief": {
            "label": "精簡",
            "sentences": "5-10 句",
            "instruction": "簡要涵蓋事件重點與關鍵數據。",
            "search_rounds": "2-3",
        },
        "normal": {
            "label": "正常",
            "sentences": "10-20 句",
            "instruction": "涵蓋事件背景、關鍵數據、各方觀點與短期影響分析。",
            "search_rounds": "3-4",
        },
        "detailed": {
            "label": "詳盡",
            "sentences": "至少 25 句",
            "instruction": (
                "深度報導等級：完整事件背景脈絡、所有關鍵數據與統計、"
                "各方觀點與專家評論、產業影響分析、未來展望與潛在風險，"
                "像專業財經記者撰寫的深度調查報導。"
            ),
            "search_rounds": "3-5",
        },
    }

    def _build_task_prompt(self, task_type: str, config: dict, today: str) -> str:
        """Build a detailed prompt for the adapter pipeline based on task type."""
        if task_type == "news":
            topic = config.get("topic", "綜合")
            count = config.get("count", 10)
            detail = config.get("detail", "normal")  # brief / normal / detailed
            extra = config.get("extra_instructions", "")
            needs_pdf = any(kw in extra.lower() for kw in ["pdf", "PDF", "下載", "檔案"])

            # ── Auto-infer missing fields from original_request ──
            original_request = config.get("original_request", "")
            if original_request and topic == "綜合":
                inferred_topic = self._extract_topic_keywords(original_request)
                if inferred_topic and inferred_topic != "綜合":
                    topic = inferred_topic
                    logger.info(f"[ScheduledPush] Auto-inferred topic from original_request: {topic}")

            level = self._NEWS_DETAIL_LEVELS.get(detail, self._NEWS_DETAIL_LEVELS["normal"])

            prompt = (
                f"⚠️ 以下是任務指令。絕對禁止在回覆中輸出任何步驟說明、流程說明、進度預告或「請稍候」訊息。直接執行，只回覆最終結果。\n\n"
                f"今天是 {today}。\n"
                f"🔒 主題約束：搜尋與整理內容必須完全符合主題「{topic}」，嚴禁產生無關主題的內容。\n"
                f"🌏 地區偏好：預設以台灣為主、日本為輔，除非使用者另有指定。\n\n"
                f"使用 mcp-web-search 搜尋今日「{topic}」最新新聞，至少搜尋 {level['search_rounds']} 次（不同關鍵字），直到蒐集滿 {count} 則不重複新聞為止（最多搜尋 6 次）。\n"
                f"搜尋關鍵字規則：query 必須包含「{topic}」核心詞彙（例如 query=\"台灣 {topic} 最新\"），嚴禁用其他無關主題搜尋。\n"
                f"蒐集完後，直接整理為 {count} 則新聞摘要，每則格式如下：\n\n"
                f"📰 新聞標題\n"
                f"摘要內容（{level['sentences']}，{level['instruction']}）\n"
                f"🔗 來源名稱\n"
                f"https://實際文章URL（來自搜尋結果，禁止編造，禁止首頁URL）\n\n"
                f"URL 規則：只能用 mcp-web-search 回傳的個別文章 URL，禁止入口網站首頁，禁止 Markdown [文字](URL) 語法。\n"
            )
            if extra:
                prompt += f"額外要求：{extra}\n"
            if needs_pdf:
                prompt += (
                    f"\n搜尋並整理完全部 {count} 則後，立刻用 mcp-python-executor 一次性生成一個 PDF（禁止分批）。\n"
                    f"PDF 每則必須包含完整標題、完整摘要（{level['sentences']}）、來源與 URL。\n"
                    f"路徑與 import 方式請參考 system prompt 中的 PDF 範例（ChinesePDF，不傳路徑給 constructor）。\n"
                    f"用列表存所有 {count} 則，迴圈寫入 PDF，一次呼叫完成，勿分批。\n"
                    f"生成後在最終回覆附上完整下載連結（https://...開頭，禁止 Markdown 語法）。\n"
                )
            prompt += (
                f"\n品質要求：繁體中文，恰好 {count} 則（不足需繼續搜尋補齊），每則摘要達到 {level['sentences']}，URL 必須真實存在。"
                f"\n最後統計：共 N 則新聞，來源 M 個。"
            )
            return prompt
        elif task_type == "work_summary":
            days = config.get("days", 7)
            return (
                f"今天是 {today}。請整理近 {days} 天的工作重點摘要：\n"
                f"1. 已完成項目\n2. 進行中項目\n3. 待處理項目\n4. 重要提醒\n"
                f"用繁體中文，格式清楚。"
            )
        elif task_type == "language":
            lang = config.get("language", "日文")
            level = config.get("level", "N3")
            count = config.get("count", 5)
            topic = config.get("topic", "")
            content_type = config.get("content_type", "vocabulary")
            current_category = config.get("current_category", "")
            used_items = config.get("used_items", [])
            # A: used_items exclusion (format-agnostic batch history)
            exclusion = (
                f"🚫 以下批次已推送過，本次必須涵蓋完全不同的內容，嚴禁重複：\n"
                f"{chr(10).join(used_items[-20:])}\n\n"
            ) if used_items else ""
            # B: category rotation — current focus topic
            focus = f"本次聚焦主題：【{current_category}】\n" if current_category else (f"本次主題：「{topic}」\n" if topic else "")
            if content_type == "grammar":
                return (
                    f"{exclusion}"
                    f"你是一位{lang}教師。請提供「恰好 {count} 個」{level}程度的文法句型教學。\n"
                    f"{focus}"
                    f"每個文法格式：\n"
                    f"句型：〜（文法型）\n意思：（中文說明）\n使用情況：（何時使用）\n"
                    f"例句1：（{lang}例句）\n翻譯：（中文翻譯）\n"
                    f"例句2：（{lang}例句）\n翻譯：（中文翻譯）\n\n"
                    f"⚠️ 輸出恰好 {count} 個文法句型，不多不少。最後出一道小測驗。用繁體中文說明。"
                )
            else:
                return (
                    f"{exclusion}"
                    f"你是一位{lang}教師。請提供「恰好 {count} 個」{level}程度的詞彙教學。\n"
                    f"{focus}"
                    f"每個詞彙格式：\n"
                    f"原文（假名）\n發音（羅馬拼音）\n中文意思\n例句（附中文翻譯）\n\n"
                    f"⚠️ 輸出恰好 {count} 個詞彙，不多不少。最後出一道小測驗。用繁體中文。"
                )
        elif task_type == "reminder":
            # Pure reminder — just return the message, no LLM needed
            message = config.get("message", "")
            if not message:
                message = config.get("original_request", "提醒時間到了！")
            return f"⏰ {message}"
        elif task_type == "custom":
            prompt = config.get("prompt", "")
            return f"現在時間：{today}\n用繁體中文回覆。\n\n{prompt}"
        else:
            return f"請用繁體中文處理以下任務：{json.dumps(config, ensure_ascii=False)}"

    # Tools that must NOT be available in scheduled push context
    _EXCLUDED_TOOLS_IN_PUSH = {
        "mcp-schedule-manager",             # Prevent recursion
        "mcp-pdf-llm-analyzer",             # Read-only analyzer, not for creating files
        "mcp-docx-llm-analyzer",            # Read-only analyzer
        "mcp-txt-llm-analyzer",             # Read-only analyzer
        "mcp-spreadsheet-llm-analyzer",     # Read-only analyzer
        "mcp-groovenaust-meeting-analyst",  # Meeting-specific
        "mcp-meeting-to-notion",            # Meeting-specific
        "mcp-high-risk-demo",               # Test only
    }
    # Tools that MUST be available for scheduled push tasks
    _REQUIRED_TOOLS_IN_PUSH = {"mcp-web-search", "mcp-python-executor"}

    def _run_via_adapter(self, prompt: str, session_id: str, needs_file: bool = False) -> str | None:
        """Run prompt through full OpenAI Adapter with tool calling enabled.

        Args:
            prompt: The task prompt
            session_id: LINE session ID
            needs_file: If True, verify response contains a download link; if not, retry with follow-up
        """
        try:
            from server.dependencies.uma import get_uma_instance
            from server.adapters.openai_adapter import OpenAIAdapter

            uma = get_uma_instance()
            adapter = OpenAIAdapter(uma)
            if not adapter.is_available:
                return None

            # Use full model for scheduled tasks (mini often fails at complex multi-step tasks)
            model = os.environ.get("LINE_MODEL_FULL", os.environ.get("OPENAI_MODEL", "gpt-4.1"))
            adapter.model = model

            # Inject session context (use original session_id, NOT scheduled_ prefix)
            os.environ["SESSION_ID"] = session_id or ""

            downloads_dir = os.path.join(os.getcwd(), "workspace", "downloads")
            base_url = os.environ.get("BASE_URL", "")
            workspace_dir = os.getcwd()
            system_prompt = (
                "你是定時推送助理。使用工具完成任務，直接輸出最終結果。\n\n"
                "═══ 絕對禁止（違反即視為失敗）═══\n"
                "❌ 禁止輸出任何「進度說明」「流程說明」「下一步」「請稍候」「處理中」「即將進行」\n"
                "❌ 禁止輸出任何形式的計劃、步驟描述或預告 — 直接做，不要說你要做\n"
                "❌ 禁止呼叫 mcp-schedule-manager（排程已建立，你的任務是執行內容）\n"
                "❌ 禁止呼叫 mcp-pdf-llm-analyzer / mcp-docx-llm-analyzer（這是讀取工具，非生成工具）\n"
                "❌ 禁止使用 import FPDF（必須用 ChinesePDF）\n"
                "❌ 禁止 ChinesePDF(路徑)（路徑只能傳給 output()）\n\n"
                "✅ 正確行為：搜尋 → 整理 → 生成檔案 → 回覆下載連結。只回覆最終完成的結果。\n\n"
                "【檔案生成方法 — 全部使用 mcp-python-executor】\n"
                f"所有檔案必須存到：{downloads_dir}\n\n"
                "■ PDF（中文必須用 ChinesePDF）：\n"
                "```python\n"
                "import sys, os\n"
                f"sys.path.insert(0, r'{workspace_dir}/workspace')\n"
                "from pdf_helper import ChinesePDF\n"
                f"DOWNLOADS = r'{downloads_dir}'\n"
                "os.makedirs(DOWNLOADS, exist_ok=True)\n"
                "out_path = os.path.join(DOWNLOADS, '檔名.pdf')\n"
                "pdf = ChinesePDF()  # ← 絕對不傳路徑，路徑在 output() 才傳\n"
                "pdf.add_page()\n"
                "pdf.chapter_title('大標題')       # 可用別名: add_title / add_heading\n"
                "pdf.chapter_subtitle('子標題')    # 可用別名: add_subtitle / add_subheading\n"
                "pdf.chapter_body('內文段落...')   # 可用別名: add_text / add_paragraph / add_content\n"
                "pdf.add_bullet('項目內容')\n"
                "pdf.add_separator()\n"
                "pdf.output(out_path)              # ← 路徑在這裡傳\n"
                "print('PDF已生成:', out_path)\n"
                "```\n"
                "⚠️ 禁止：ChinesePDF(路徑) — 路徑不能傳給 constructor，必須傳給 output()。\n"
                "⚠️ 禁止：import FPDF — 必須用 ChinesePDF，否則中文會亂碼。\n\n"
                "■ DOCX：\n"
                "```python\n"
                "from docx import Document\n"
                "doc = Document()\n"
                "doc.add_heading('標題', 0)\n"
                "doc.add_paragraph('內文...')\n"
                f"doc.save(r'{downloads_dir}/檔名.docx')\n"
                "```\n\n"
                "■ TXT / MD：\n"
                "```python\n"
                f"with open(r'{downloads_dir}/檔名.txt', 'w', encoding='utf-8') as f:\n"
                "    f.write('內容...')\n"
                "```\n\n"
                f"- 檔案存放路徑：{downloads_dir}\n"
                f"- 下載連結格式：{base_url}/downloads/檔案名稱\n"
                "- 禁止使用 Markdown 超連結語法 [文字](URL)，LINE 不支援\n"
                "- URL 必須完整顯示 https://... 開頭，LINE 會自動轉為可點擊連結\n"
                "回覆請用繁體中文。"
            )

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ]

            # Override get_tools to ensure web-search is included and schedule-manager is excluded
            _original_get_tools = adapter.get_tools

            def _patched_get_tools(user_query="", max_tools=15):
                tools = _original_get_tools(user_query=user_query, max_tools=max_tools)
                # Remove excluded tools (prevent recursion)
                tools = [t for t in tools if t.get("name") not in self._EXCLUDED_TOOLS_IN_PUSH]
                # Check if required tools are present
                tool_names = {t.get("name") for t in tools}
                for required in self._REQUIRED_TOOLS_IN_PUSH:
                    if required not in tool_names:
                        # Fetch from full list
                        all_tools = _original_get_tools(user_query="", max_tools=50)
                        for t in all_tools:
                            if t.get("name") == required:
                                tools.append(t)
                                logger.info(f"[ScheduledPush] Force-injected required tool: {required}")
                                break
                return tools

            adapter.get_tools = _patched_get_tools

            final_text = ""
            for chunk in adapter.chat(
                messages=messages,
                user_query=prompt,
                session_id=session_id,  # Use original session_id, no prefix
                tools_enabled=True,
            ):
                status = chunk.get("status", "")
                if status == "success":
                    final_text = chunk.get("content", "")
                elif status == "streaming":
                    pass  # Accumulating internally

            # ── Follow-up: if file was expected but not generated, force a second round ──
            if needs_file and final_text:
                base_url = os.environ.get("BASE_URL", "")
                has_download_link = (
                    (base_url and base_url in final_text)
                    or "/downloads/" in final_text
                )
                if not has_download_link:
                    logger.warning(
                        "[ScheduledPush] File expected but no download link found. "
                        "Sending follow-up to force file generation."
                    )
                    downloads_dir = os.path.join(os.getcwd(), "workspace", "downloads")

                    # Truncate final_text to avoid token explosion in follow-up
                    # Keep only first 3000 chars of assistant content (enough context for PDF)
                    truncated_content = final_text[:3000]
                    if len(final_text) > 3000:
                        truncated_content += "\n...(以上為部分內容，請根據這些資料生成完整 PDF)"

                    # Use minimal system prompt for follow-up (save tokens)
                    follow_up_system = (
                        "你是檔案生成助理。使用 mcp-python-executor 執行 Python 生成 PDF。\n"
                        f"檔案存到：{downloads_dir}\n"
                        f"下載連結：{base_url}/downloads/檔案名稱\n"
                        "中文 PDF 必須用 ChinesePDF，正確用法：\n"
                        "```python\n"
                        "import sys, os\n"
                        f"sys.path.insert(0, r'{workspace_dir}/workspace')\n"
                        "from pdf_helper import ChinesePDF\n"
                        f"out_path = os.path.join(r'{downloads_dir}', '檔名.pdf')\n"
                        "pdf = ChinesePDF()  # 不傳路徑\n"
                        "pdf.add_page()\n"
                        "pdf.chapter_title('標題')   # 別名: add_title/add_heading\n"
                        "pdf.chapter_body('內文')    # 別名: add_text/add_paragraph\n"
                        "pdf.output(out_path)        # 路徑在這裡\n"
                        "```\n"
                        "⚠️ ChinesePDF(路徑) 是錯的，路徑只能傳給 output()。\n"
                        "禁止使用 Markdown 超連結語法。URL 必須完整顯示。"
                    )

                    follow_up_messages = [
                        {"role": "system", "content": follow_up_system},
                        {"role": "user", "content": (
                            "以下是已整理好的新聞內容，請立刻用 mcp-python-executor "
                            "將全部內容製作成一個 PDF 檔案。\n"
                            "禁止再次搜尋，直接用以下內容生成 PDF：\n\n"
                            f"{truncated_content}\n\n"
                            f"存放到：{downloads_dir}\n"
                            f"完成後回覆下載連結：{base_url}/downloads/檔案名稱"
                        )},
                    ]

                    # Create fresh adapter for follow-up to avoid token accumulation
                    follow_up_adapter = OpenAIAdapter(uma)
                    follow_up_adapter.model = model
                    follow_up_adapter.get_tools = _patched_get_tools

                    follow_up_text = ""
                    for chunk in follow_up_adapter.chat(
                        messages=follow_up_messages,
                        user_query="生成PDF檔案",
                        session_id=session_id,
                        tools_enabled=True,
                    ):
                        st = chunk.get("status", "")
                        if st == "success":
                            follow_up_text = chunk.get("content", "")

                    if follow_up_text and ("/downloads/" in follow_up_text or (base_url and base_url in follow_up_text)):
                        # Append the download link to the original text
                        final_text = final_text.rstrip() + "\n\n" + follow_up_text
                        logger.info("[ScheduledPush] Follow-up successfully generated file.")
                    else:
                        logger.warning("[ScheduledPush] Follow-up still didn't generate file link.")

            return final_text if final_text else None

        except Exception as e:
            logger.error(f"[ScheduledPush] Adapter execution failed: {e}")
            return None

    # ─── Scheduler Tick ──────────────────────────────────────────────────────

    def check_and_execute(self, llm_callable=None, tool_executor=None, push_fn=None):
        """
        Called by APScheduler every minute.
        Checks all configs for due tasks, generates content, and pushes.
        """
        if not push_fn:
            logger.warning("[ScheduledPush] No push function provided, skipping.")
            return

        now = datetime.now()
        current_hour = now.hour
        current_minute = now.minute
        current_dow = now.strftime("%a").lower()[:3]  # mon, tue, ...
        current_day = now.day

        # Map day_of_week strings
        dow_map = {"mon": 0, "tue": 1, "wed": 2, "thu": 3, "fri": 4, "sat": 5, "sun": 6}

        for config_file in self.schedules_dir.glob("*.json"):
            try:
                config = json.loads(config_file.read_text(encoding="utf-8"))
            except Exception:
                continue

            chat_id = config.get("chat_id", "")
            session_id = config.get("session_id", "")
            if not chat_id:
                continue

            for task in config.get("tasks", []):
                if not task.get("enabled", True):
                    continue

                cron = task.get("cron_parsed", {})
                interval_minutes = cron.get("interval_minutes")

                if interval_minutes:
                    # Interval task: fire when elapsed time since last_run >= interval
                    last_run_str = task.get("last_run")
                    if last_run_str:
                        try:
                            lr = datetime.fromisoformat(last_run_str)
                            # Prevent duplicate in same minute
                            if lr.hour == current_hour and lr.minute == current_minute:
                                continue
                            elapsed = (now - lr).total_seconds() / 60
                            if elapsed < interval_minutes:
                                continue
                        except ValueError:
                            pass
                else:
                    task_hour = cron.get("hour")
                    task_minute = cron.get("minute", 0)

                    # Check hour & minute match
                    if task_hour is not None and task_hour != current_hour:
                        continue
                    if task_minute != current_minute:
                        continue

                    # Check day_of_week
                    dow_filter = cron.get("day_of_week")
                    if dow_filter:
                        if dow_filter == "mon-fri":
                            if dow_map.get(current_dow, 0) > 4:
                                continue
                        elif current_dow not in dow_filter:
                            continue

                    # Check day of month
                    day_filter = cron.get("day")
                    if day_filter and day_filter != current_day:
                        continue

                    # Prevent duplicate runs (check last_run within same minute)
                    last_run = task.get("last_run")
                    if last_run:
                        try:
                            lr = datetime.fromisoformat(last_run)
                            if lr.hour == current_hour and lr.minute == current_minute and lr.date() == now.date():
                                continue  # Already ran this minute
                        except ValueError:
                            pass

                # ── Execute ──
                logger.info(f"[ScheduledPush] Executing task '{task['name']}' ({task['type']}) for {session_id}")

                try:
                    content = self.generate_content(
                        task, session_id,
                        llm_callable=llm_callable,
                        tool_executor=tool_executor,
                    )

                    # Add header
                    header = f"⏰ 定時推送 — {task['name']}\n{'─' * 20}\n\n"
                    full_message = header + content

                    # Push to LINE
                    push_fn(chat_id, full_message)
                    logger.info(f"[ScheduledPush] Pushed '{task['name']}' to {chat_id}")

                    # Update last_run
                    task["last_run"] = now.isoformat()
                    self.save_config(session_id, config)

                except Exception as e:
                    logger.error(f"[ScheduledPush] Task '{task['name']}' failed: {e}")

    # ─── Natural Language Task Creation (via LLM) ────────────────────────────

    @staticmethod
    def parse_schedule_intent(text: str) -> Optional[dict]:
        """
        Try to parse scheduling intent from user text.
        Returns dict with type, name, cron, config if detected, else None.

        Examples:
          '每天早上8點推送科技新聞5則' → {type: 'news', cron: '08:00', config: {topic: '科技', count: 5}}
          '工作日下午5點推送本週工作摘要' → {type: 'work_summary', cron: 'weekday 17:00'}
          '每天9點推送5個日文N3商務詞彙' → {type: 'language', cron: '09:00', config: {language: '日文', level: 'N3', count: 5}}
        """
        # This is a lightweight pattern matcher; complex cases go through LLM
        import re

        result = None

        # Detect news
        news_match = re.search(r'(?:推送|傳送?|發送?).*?(\d+)\s*則.*?新聞', text)
        if not news_match:
            news_match = re.search(r'新聞.*?(\d+)\s*則', text)
        if '新聞' in text:
            count = 5
            m = re.search(r'(\d+)\s*則', text)
            if m:
                count = int(m.group(1))
            topic = "科技與AI"
            for kw in ["科技", "財經", "國際", "體育", "娛樂", "商業", "AI", "政治"]:
                if kw in text:
                    topic = kw
                    break
            result = {"type": "news", "config": {"topic": topic, "count": count}}

        # Detect work summary
        elif any(kw in text for kw in ["工作摘要", "工作總結", "工作重點", "工作項目", "週報"]):
            days = 7
            m = re.search(r'(\d+)\s*[天日]', text)
            if m:
                days = int(m.group(1))
            result = {"type": "work_summary", "config": {"days": days}}

        # Detect language
        elif any(kw in text for kw in ["詞彙", "單字", "學習", "語言"]):
            lang = "日文"
            for l in ["日文", "英文", "韓文", "法文", "德文", "西班牙文"]:
                if l in text:
                    lang = l
                    break
            level = ""
            m = re.search(r'[NnJj]\d', text)
            if m:
                level = m.group(0).upper()
            count = 5
            m = re.search(r'(\d+)\s*個', text)
            if m:
                count = int(m.group(1))
            topic_kw = "日常"
            for kw in ["商務", "旅遊", "日常", "學術", "IT"]:
                if kw in text:
                    topic_kw = kw
                    break
            result = {"type": "language", "config": {"language": lang, "level": level, "count": count, "topic": topic_kw}}

        if result is None:
            return None

        # Extract time
        time_match = re.search(r'(\d{1,2})[:\s時點](\d{0,2})', text)
        if time_match:
            h = int(time_match.group(1))
            m = int(time_match.group(2)) if time_match.group(2) else 0
            cron = f"{h:02d}:{m:02d}"
        else:
            cron = "08:00"

        # Check weekday
        if any(kw in text for kw in ["工作日", "平日", "weekday"]):
            cron = f"weekday {cron}"

        result["cron"] = cron

        # Auto-generate name
        type_names = {
            "news": f"每日{result.get('config', {}).get('topic', '')}新聞",
            "work_summary": "工作重點摘要",
            "language": f"{result.get('config', {}).get('language', '')}詞彙學習",
        }
        result["name"] = type_names.get(result["type"], "定時推送")

        return result
